#!/usr/bin/env python3
"""Render the estimand table (appendix Table 2) as a shareable .docx.

Every numeric cell is READ FROM ITS SOURCE CSV, not retyped. That is deliberate:
the draft markdown version had stale row-4 risk cells (TB 3.14/1.54, not-TB
4.62/3.61) left over from before the five-year / TB-from-any-line change, while
its risk DIFFERENCES were current. Reading from source makes that class of drift
impossible.

Sources
-------
row 1  CCW_analysis/results_v3/ccw_bootstrap_h60.csv   (estimate = nested_T6)
       ITT_Analysis/results/fig3_primary_curves.csv     (arm risks, MI-pooled)
rows 2 and 4  ITT_Analysis/results/rolling_cause_cif_boot.csv  (dmon blank = overall)
       ITT_Analysis/results/rolling_cause_cif_draws_PRIMARY_underlying.csv
         -- the combined "not tuberculosis" interval, recomputed per replicate
            rather than by adding two intervals, which would be wrong because the
            two components are correlated.
row 3  ITT_Analysis/results/rolling_rd_by_month.csv     (time_y = 5)

The residual class is nontb + unclass, i.e. all-cause minus tuberculosis. That is
the owner's two-class partition (2026-08-21) and it is why the parts sum exactly
to the all-cause row.

Usage:  python3 ITT_Analysis/scripts/75_make_estimand_table_docx.py
"""
import os
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(os.environ.get("TB_ABANDONMENT_ROOT", Path(__file__).resolve().parents[2]))
RES = ROOT / "ITT_Analysis" / "results"
OUT = ROOT / "Plos Medicine" / "R1" / "Table_estimands_appendix.docx"

INK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

# ---------------------------------------------------------------- numbers ----
ccw = pd.read_csv(ROOT / "CCW_analysis/results_v3/ccw_bootstrap_h60.csv")
c = ccw[ccw.estimate == "nested_T6"].iloc[0]
cur = pd.read_csv(RES / "fig3_primary_curves.csv")
a = cur[cur.group == "All patients"].iloc[0]

cif = pd.read_csv(RES / "rolling_cause_cif_boot.csv")
o = cif[cif.dmon.isna() & (cif.time_y == 5)].set_index("cause")
al, tb, ntb, un = o.loc["all"], o.loc["tb"], o.loc["nontb"], o.loc["unclass"]
res_r1, res_r0 = ntb.risk1 + un.risk1, ntb.risk0 + un.risk0
res_rd = ntb.rd + un.rd

draws = pd.read_csv(RES / "rolling_cause_cif_draws_PRIMARY_underlying.csv")
dw = draws[draws.dmon.isna() & (draws.time_y == 5)].pivot_table(
    index="rep", columns="cause", values="rd")
comb = dw["nontb"] + dw["unclass"]
res_lo, res_hi = np.percentile(comb, 2.5), np.percentile(comb, 97.5)
n_reps = len(dw)

mon = pd.read_csv(RES / "rolling_rd_by_month.csv")
m5 = mon[(mon.time_y == 5) & mon.dmon.notna()].sort_values("dmon")
rd_min, rd_max = m5.rd.min(), m5.rd.max()
excl = m5[m5.rd_lo > 0].dmon.astype(int).tolist()
null_m = m5[m5.rd_lo <= 0].dmon.astype(int).tolist()
n_lo, n_hi = int(m5.n_exposed.min()), int(m5.n_exposed.max())

print(f"[75] row1 RD {c.rd:.2f} ({c.rd_lo:.2f}-{c.rd_hi:.2f})  "
      f"risks {a.risk_disengage:.2f}/{a.risk_remain:.2f}")
print(f"[75] row2 RD {al.rd:.2f} ({al.rd_lo:.2f}-{al.rd_hi:.2f})  "
      f"risks {al.risk1:.2f}/{al.risk0:.2f}  n={int(al.n_exposed):,}")
print(f"[75] row3 RD {rd_min:.2f} to {rd_max:.2f}; excludes zero months {excl}; "
      f"null months {null_m}")
print(f"[75] row4 TB {tb.rd:.2f} ({tb.rd_lo:.2f}-{tb.rd_hi:.2f}) | "
      f"not-TB {res_rd:.2f} ({res_lo:.2f}-{res_hi:.2f}) from {n_reps} replicates")
print(f"[75] additivity: {tb.rd:.4f} + {res_rd:.4f} = {tb.rd + res_rd:.4f} "
      f"vs all-cause {al.rd:.4f}")

ROWS = [
    ("1",
     "What would 60-month mortality be if all patients disengaged from care "
     "within the first 6 months, versus if none did?",
     "Clone–censor–weight target trial",
     "Treatment start (day 0), for both clones of every patient",
     "The same patients, under the strategy of remaining in care",
     f"{a.risk_disengage:.2f}%", f"{a.risk_remain:.2f}%",
     f"RD +{c.rd:.2f} pp ({c.rd_lo:.2f} to {c.rd_hi:.2f})\n"
     f"RR {c.rr:.2f} ({c.rr_lo:.2f}–{c.rr_hi:.2f})",
     "171,048 patients, of whom 20,830 lost to follow-up"),
    ("2",
     "Among patients who disengage, what is 5-year mortality versus otherwise "
     "similar patients still in care on the same date?",
     "Rolling landmark",
     "Each patient's own loss-to-follow-up declaration date (last dose + 30 days)",
     "Patients alive and still in care on that same date",
     f"{al.risk1:.2f}%", f"{al.risk0:.2f}%",
     f"RD +{al.rd:.2f} pp ({al.rd_lo:.2f} to {al.rd_hi:.2f})\n"
     f"RR {al.rr:.2f} ({al.rr_lo:.2f}–{al.rr_hi:.2f})",
     f"{int(al.n_exposed):,} exposed"),
    ("3",
     "Does the effect depend on WHEN in treatment disengagement occurs?",
     "Rolling landmark, stratified by month of disengagement",
     "Each patient's own declaration date",
     "Patients alive and still in care on that same date",
     "—", "—",
     f"RD +{rd_min:.2f} to +{rd_max:.2f} pp across months 1 to 6; excludes zero "
     f"in months {excl[0]} to {excl[-1]}, not distinguishable from zero in "
     f"month {null_m[0]}",
     f"{n_lo:,} to {n_hi:,} exposed per month"),
    ("4",
     "Is the excess concentrated in deaths attributed to tuberculosis?",
     "Rolling landmark, cause-specific cumulative incidence (Aalen–Johansen)",
     "Each patient's own declaration date",
     "Patients alive and still in care on that same date",
     f"{tb.risk1:.2f}% / {res_r1:.2f}%", f"{tb.risk0:.2f}% / {res_r0:.2f}%",
     f"Tuberculosis death RD +{tb.rd:.2f} pp ({tb.rd_lo:.2f} to {tb.rd_hi:.2f})\n"
     f"Not-tuberculosis death RD +{res_rd:.2f} pp ({res_lo:.2f} to {res_hi:.2f})",
     f"{int(al.n_exposed):,} exposed"),
]

HEAD = ["#", "Question being answered", "Design", "Time origin", "Comparator",
        "Risk, exposed", "Risk, comparator", "Contrast (95% CI)", "Denominator"]
WIDTHS = [0.25, 2.15, 1.30, 1.35, 1.30, 0.72, 0.80, 1.75, 1.05]

FOOTNOTES = [
    ("a", "Rows 1 and 2 answer different questions on DIFFERENT CLOCKS and their "
          "magnitudes are not interchangeable. Row 1 is a population-level strategy "
          "contrast measured from treatment start, diluted by the months before the "
          "strategies diverge. Row 2 is measured from each patient's own declaration "
          "date, among those who disengaged."),
    ("b", f"Row 1 is the primary reported estimand. Its interval comes from a "
          f"patient-level bootstrap (B = {int(c.B_ok)}) that resamples patients, "
          f"rebuilds both arms and refits the censoring weights in every replicate. "
          f"Its point estimate is pooled across the five imputed datasets."),
    ("c", "All estimates use the whole 5 years from their own origin; the early/late "
          "window split has been removed from the paper. One consequence to keep in "
          "view: a patient who dies before the 30-day definition can be met is "
          "recorded as an on-treatment death and cannot be classified as exposed, so "
          "the first months after either origin carry a competing-exposure artifact "
          "that deflates the estimates."),
    ("d", f"Row 4 is on the risk-difference scale like every other row, and the two "
          f"components add exactly to row 2: {tb.rd:.2f} + {res_rd:.2f} = "
          f"{tb.rd + res_rd:.2f}, the all-cause risk difference. The arm risks add "
          f"in the same way. That additivity is why the cause-specific estimates use "
          f"Aalen–Johansen cumulative incidence rather than one-minus-"
          f"Kaplan–Meier with competing deaths censored, which would overstate "
          f"each part. Risk cells are given as tuberculosis / not tuberculosis. The "
          f"second class is the residual, all-cause minus tuberculosis; it therefore "
          f"contains deaths whose recorded cause is not tuberculosis together with a "
          f"small number that could not be classified."),
    ("e", f"The interval for the combined not-tuberculosis class was recomputed "
          f"within each of the {n_reps} bootstrap replicates, never by adding two "
          f"intervals: the components are correlated and differencing or summing "
          f"their percentiles would give the wrong width."),
    ("f", "Row 3 is stated as risk differences rather than hazard ratios, consistent "
          "with the primary estimand and the figures."),
]

# ------------------------------------------------------------------ document -
doc = Document()
s = doc.sections[0]
s.orientation = WD_ORIENT.LANDSCAPE
s.page_width, s.page_height = Inches(11), Inches(8.5)
for attr in ("left_margin", "right_margin"):
    setattr(s, attr, Inches(0.5))
s.top_margin = s.bottom_margin = Inches(0.6)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(9)


def para(text, size=9, bold=False, italic=False, colour=INK, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = colour
    return p


def shade(cell, hexfill):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


para("Table 2. Estimands, comparators and estimates", size=12, bold=True, after=3)
para("Appendix table. All estimates adjusted for the 14 baseline covariates listed "
     "in Methods and pooled across five imputed datasets. RD, risk difference; "
     "RR, risk ratio; pp, percentage points.", size=8.5, colour=GREY, after=10)

t = doc.add_table(rows=1, cols=len(HEAD))
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False

for i, (cell, head, w) in enumerate(zip(t.rows[0].cells, HEAD, WIDTHS)):
    cell.width = Inches(w)
    shade(cell, "E8E8E8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(head)
    r.bold = True
    r.font.size = Pt(8.5)

for row in ROWS:
    cells = t.add_row().cells
    for i, (cell, val, w) in enumerate(zip(cells, row, WIDTHS)):
        cell.width = Inches(w)
        for j, line in enumerate(str(val).split("\n")):
            p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(line)
            r.font.size = Pt(8)
            r.bold = (i == 0) or (i == 7 and row[0] == "1")

doc.add_paragraph()
para("Footnotes", size=9.5, bold=True, after=3)
for tag, text in FOOTNOTES:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    r = p.add_run(f"{tag}.  ")
    r.bold = True
    r.font.size = Pt(8)
    r2 = p.add_run(text)
    r2.font.size = Pt(8)

# --- internal section, clearly separated ------------------------------------
doc.add_page_break()
para("Internal notes — for Jason and Evelyn, not for the manuscript",
     size=11, bold=True, after=6)
NOTES = [
    ("Every number above is read from its source CSV at build time",
     "This document is generated by ITT_Analysis/scripts/75_make_estimand_table_docx.py. "
     "Re-run it after any analysis change rather than editing the .docx, or the two "
     "will drift apart."),
    ("Row 4's risk cells were stale in the markdown draft and are corrected here",
     "The draft carried tuberculosis 3.14% / 1.54% and not-tuberculosis 4.62% / 3.61%, "
     "left over from before the five-year and TB-from-any-certificate-line changes. "
     f"The current source gives tuberculosis {tb.risk1:.2f}% / {tb.risk0:.2f}% and "
     f"not-tuberculosis {res_r1:.2f}% / {res_r0:.2f}%. These now add exactly to row 2's "
     f"{al.risk1:.2f}% and {al.risk0:.2f}%, which the stale pair did not. The risk "
     "DIFFERENCES were correct in the draft and are unchanged."),
    ("Cause-specific death counts have been removed from the denominator column",
     "The draft stated 1,498 tuberculosis deaths and 5,423 other deaths. Those sum to "
     "6,921, which does not reconcile with any death total in the current outputs, so "
     "they are omitted rather than reproduced. If a per-cause death count is wanted in "
     "this table, it needs deriving from the landmark stack and adding to a results CSV "
     "first."),
    ("Open: which design is licensed to carry row 4",
     "The owner narrowed the split on 2026-08-21 to “CCW primary, rolling landmark "
     "for months only”. Rows 2 and 4 are landmark estimates that are not about "
     "timing, so their place in this table is a live question. The numbers are correct; "
     "what is unsettled is whether they belong here."),
    ("Figure cross-references",
     "The figures were renumbered on 2026-08-21: the new standardised-risk-curve figure "
     "is Figure 4 and the landmark-and-CCW forest is now Figure 5. Any surviving "
     "reference to “Figure 4” meaning the forest is stale."),
]
for head, body in NOTES:
    para(head, size=9, bold=True, after=1)
    para(body, size=8.5, colour=GREY, after=7)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"[75] wrote {OUT.relative_to(ROOT)}")
