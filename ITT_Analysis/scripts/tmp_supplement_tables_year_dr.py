"""
Build supplement tables for the year-trends and drug-resistance sensitivity
analyses. Produces three .docx tables in ITT_Analysis/results/Tables/ to
match existing Table1_ITT.docx / Table2_ITT.docx style:

  S_year_ltfu.docx        LTFU rate by year of treatment start, 2013-2023
  S_year_outcomes_ltfu.docx 1y / 2y mortality + retreatment among LTFU,
                            by year of treatment start, 2013-2022
  S_dr_status.docx        Drug-resistance status (Sensitive / INH-mono /
                            RR-MDR / Not evaluated), overall and by itt_group

Inputs (already produced by earlier exploratory scripts):
  results/year_trends/ltfu_by_year.csv
  results/year_trends/outcomes_ltfu_by_year.csv
  results/resistance/dr_status_lookup.csv
  ITT_Analysis/data/itt_cohort.csv
"""

import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from scipy import stats


def _find_project_root() -> Path:
    if os.environ.get("TB_ABANDONMENT_ROOT"):
        p = Path(os.environ["TB_ABANDONMENT_ROOT"]).expanduser()
        if p.exists():
            return p
    for c in [
        Path.home() / "Library/CloudStorage/GoogleDrive-jasonandr@gmail.com/My Drive/Abandonment Paper",
        Path.home() / "Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/Abandonment Outcomes/Abandonment Paper",
    ]:
        if c.exists():
            return c
    return Path(__file__).resolve().parents[2]


BASE = _find_project_root()
RES = BASE / "ITT_Analysis" / "results"
OUT = RES / "Tables"
OUT.mkdir(parents=True, exist_ok=True)


def _bold_header(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def _add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _build_doc(title, headers, rows, note=None):
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    _bold_header(table.rows[0])
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = str(val)
    if note:
        _add_note(doc, note)
    return doc


# ---------------------------------------------------------------------------
# S1: LTFU rate by year of treatment start, 2013-2023
# ---------------------------------------------------------------------------
ltfu = pd.read_csv(RES / "year_trends" / "ltfu_by_year.csv")
headers_s1 = ["Year of treatment start", "Total cohort, N",
              "Non-LTFU, n", "LTFU, n", "LTFU % (95% CI)"]
rows_s1 = []
for _, r in ltfu.iterrows():
    yr = int(r["year_start"])
    label = f"{yr}*" if yr == 2023 else str(yr)
    rows_s1.append([
        label,
        f"{int(r['n_total']):,}",
        f"{int(r['n_non_ltfu']):,}",
        f"{int(r['n_ltfu']):,}",
        f"{r['ltfu_pct']:.1f} ({r['ltfu_ci_lo']:.1f}–{r['ltfu_ci_hi']:.1f})",
    ])
note_s1 = (
    "LTFU is ascertained at index-episode end (case_outcome ∈ "
    "{Abandono, Abandono Primario, Faltoso}). 95% confidence intervals are "
    "Wilson score intervals. *2023 cohort is right-truncated by the inclusion "
    "criterion end_date ≤ 2023-12-31: late-2023 starters whose treatment ends "
    "in 2024 are excluded, disproportionately removing Non-LTFU (longer "
    "treatments). The 2023 LTFU% is therefore upward-biased and is not "
    "directly comparable with prior years."
)
title_s1 = "Table S_year_ltfu. Loss to follow-up at index episode, by year of treatment start (São Paulo, 2013–2023)"
doc_s1 = _build_doc(title_s1, headers_s1, rows_s1, note=note_s1)
out_s1 = OUT / "S_year_ltfu.docx"
doc_s1.save(out_s1)
ltfu.to_csv(OUT / "S_year_ltfu.csv", index=False)
print(f"Wrote {out_s1}")


# ---------------------------------------------------------------------------
# S2: 1y/2y mortality and retreatment among LTFU, by year, 2013-2022
# ---------------------------------------------------------------------------
out_l = pd.read_csv(RES / "year_trends" / "outcomes_ltfu_by_year.csv")
headers_s2 = [
    "Year of treatment start",
    "LTFU, N",
    "1-year mortality % (95% CI)",
    "2-year mortality % (95% CI)",
    "1-year retreatment % (95% CI)",
    "2-year retreatment % (95% CI)",
]
rows_s2 = []
for _, r in out_l.iterrows():
    rows_s2.append([
        int(r["year_start"]),
        f"{int(r['n_ltfu']):,}",
        f"{r['mort_1y_pct']:.1f} ({r['mort_1y_lo']:.1f}–{r['mort_1y_hi']:.1f})",
        f"{r['mort_2y_pct']:.1f} ({r['mort_2y_lo']:.1f}–{r['mort_2y_hi']:.1f})",
        f"{r['retreat_1y_pct']:.1f} ({r['retreat_1y_lo']:.1f}–{r['retreat_1y_hi']:.1f})",
        f"{r['retreat_2y_pct']:.1f} ({r['retreat_2y_lo']:.1f}–{r['retreat_2y_hi']:.1f})",
    ])
note_s2 = (
    "Restricted to patients with itt_group = Loss to follow-up. Time origin "
    "is treatment start (best_start). Mortality at K years = death within K "
    "years of treatment start (binomial proportion; Wilson 95% CI). "
    "Retreatment at K years = re-notification within K years, with death as "
    "a competing event; reported as the cumulative incidence at K years "
    "(numerically equivalent to a simple proportion here because all years "
    "shown have ≥K years of potential follow-up given a 2024-12-31 "
    "administrative censor). 2023 starters are excluded because the cohort "
    "is right-truncated at end_date ≤ 2023-12-31, biasing the 2023 LTFU "
    "subgroup."
)
title_s2 = "Table S_year_outcomes. One- and two-year mortality and retreatment among LTFU patients, by year of treatment start (2013–2022)"
doc_s2 = _build_doc(title_s2, headers_s2, rows_s2, note=note_s2)
out_s2 = OUT / "S_year_outcomes_ltfu.docx"
doc_s2.save(out_s2)
out_l.to_csv(OUT / "S_year_outcomes_ltfu.csv", index=False)
print(f"Wrote {out_s2}")


# ---------------------------------------------------------------------------
# S3: dr_status overall and by itt_group, with LTFU% per category
# ---------------------------------------------------------------------------
cohort = pd.read_csv(
    BASE / "ITT_Analysis" / "data" / "itt_cohort.csv",
    usecols=["sinan_clean", "itt_group"],
    low_memory=False,
)
dr = pd.read_csv(RES / "resistance" / "dr_status_lookup.csv")
m = cohort.merge(dr, on="sinan_clean", how="left")
m["dr_status"] = m["dr_status"].fillna("Not Evaluated")

cat_order = ["RR/MDR-TB", "INH-mono resistance", "Sensitive", "Not Evaluated"]

n_total = len(m)
n_ltfu_all = (m["itt_group"] == "Loss to follow-up").sum()
n_non_all = (m["itt_group"] == "Non-LTFU").sum()

headers_s3 = [
    "Drug-resistance status",
    f"Total (N={n_total:,}), n (%)",
    f"Non-LTFU (N={n_non_all:,}), n (%)",
    f"LTFU (N={n_ltfu_all:,}), n (%)",
    "LTFU % within row (95% CI)",
]
rows_s3 = []
for cat in cat_order:
    sub = m[m["dr_status"] == cat]
    n = len(sub)
    n_ltfu = (sub["itt_group"] == "Loss to follow-up").sum()
    n_non = (sub["itt_group"] == "Non-LTFU").sum()
    p = n_ltfu / n if n else 0.0
    if n > 0:
        lo, hi = stats.binomtest(int(n_ltfu), n).proportion_ci(method="wilson")
    else:
        lo = hi = 0.0
    rows_s3.append([
        cat,
        f"{n:,} ({n / n_total * 100:.1f}%)",
        f"{n_non:,} ({n_non / n_non_all * 100:.1f}%)",
        f"{n_ltfu:,} ({n_ltfu / n_ltfu_all * 100:.1f}%)",
        f"{p * 100:.1f} ({lo * 100:.1f}–{hi * 100:.1f})",
    ])

note_s3 = (
    "Drug-resistance status was derived per patient by rolling up four raw "
    "SINAN variables across all notifications (tmr_tb [Xpert MTB/RIF], "
    "resistance [DST overall], rifasens [rifampin DST], isonisens [isoniazid "
    "DST]) using a hierarchical rule: (1) RR/MDR-TB if any direct evidence of "
    "rifampin resistance (tmr_tb = RIF-resistant, rifasens = Resist, or "
    "resistance = TB MR); else (2) INH-mono resistance if isonisens = Resist; "
    "else (3) Sensitive if any positive evidence of sensitivity; else (4) "
    "Not Evaluated. Patients with resistance = TB R but no specific "
    "drug-level Resist call (n=9 in the cohort) were classified as Not "
    "Evaluated. Patients with mixed Sens/Resist results across rows (n=588 "
    "in the raw file, ≈0.3%) were classified by their resistance signal. "
    "The Not Evaluated category is large (~58%) and likely reflects a "
    "non-random subgroup (more often smear-negative or extrapulmonary "
    "disease, where DST is less commonly performed)."
)
title_s3 = "Table S_dr_status. Drug-resistance status by treatment-outcome group (ITT cohort, N=172,463)"
doc_s3 = _build_doc(title_s3, headers_s3, rows_s3, note=note_s3)
out_s3 = OUT / "S_dr_status.docx"
doc_s3.save(out_s3)
pd.DataFrame(rows_s3, columns=headers_s3).to_csv(OUT / "S_dr_status.csv", index=False)
print(f"Wrote {out_s3}")

print("Done.")
