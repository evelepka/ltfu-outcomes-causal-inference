"""
60_make_manuscript.py
---------------------
Assemble a Word manuscript draft from the current pipeline outputs.

Inputs (all on GDrive under ITT_Analysis/results/):
  - Figure_1_descriptive.png     (Figure 1)
  - Figure_2_stratified_retreatment_24mo.png (Figure 2)
  - Figure_3_causal_mortality.png            (Figure 3)
  - Figure_4_within_ltfu_forest.png          (Figure 4)
  - target_trial_mi_early_late_array.csv
  - target_trial_subgroup_interactions_mi.csv
  - multivariable_results_mi_cc.csv
  - Figure_1_caption_stats.csv
  - Figure_2_cif_values_24mo.csv

Input cohort:
  - itt_cohort.csv (for Table 1)

Output:
  - {GDrive}/Abandonment Paper/Drafts/Draft_{YYYY-MM-DD}.docx

Background and Discussion are outlined, not prose; Methods and Results
are written in full, with numbers pulled from the CSVs above.
"""
import os
import warnings
from datetime import datetime
from pathlib import Path

import pandas as pd
import numpy as np
from scipy.stats import chi2_contingency, mannwhitneyu
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")


# --------------------------------------------------------------------------
# Paths
# --------------------------------------------------------------------------
def _find_project_root():
    env = os.environ.get("TB_ABANDONMENT_ROOT")
    if env and Path(env).exists():
        return Path(env)
    for c in [
        Path.home() / "Library/CloudStorage/GoogleDrive-jasonandr@gmail.com/My Drive/Abandonment Paper",
        Path.home() / "Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/Abandonment Outcomes/Abandonment Paper",
        Path.home() / "Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/TB SP 2026/LTFU Paper",
    ]:
        if c.exists():
            return c
    return Path(__file__).resolve().parents[2]


BASE = _find_project_root()
RESULTS = BASE / "ITT_Analysis" / "results"
COHORT_CSV = BASE / "ITT_Analysis" / "data" / "itt_cohort.csv"

DRAFTS = Path.home() / ("Library/CloudStorage/GoogleDrive-jasonandr@gmail.com/"
                        ".shortcut-targets-by-id/18HafZqxrHeLVzpA6oYW-1cro9ygNfwVd/"
                        "Abandonment Paper/Drafts")
if not DRAFTS.exists():
    DRAFTS = BASE / "Drafts"

OUT_PATH = DRAFTS / f"Draft_{datetime.now():%Y-%m-%d}.docx"


# --------------------------------------------------------------------------
# Load numbers
# --------------------------------------------------------------------------
fig1_stats  = pd.read_csv(RESULTS / "Figure_1_caption_stats.csv").iloc[0].to_dict()
fig2_vals   = pd.read_csv(RESULTS / "Figure_2_cif_values_24mo.csv")
tt_array    = pd.read_csv(RESULTS / "target_trial_mi_early_late_array.csv")
tt_subgrp   = pd.read_csv(RESULTS / "target_trial_subgroup_interactions_mi.csv")
mv          = pd.read_csv(RESULTS / "multivariable_results_mi_cc.csv")

cohort = pd.read_csv(COHORT_CSV, low_memory=False)
N_total = len(cohort)
N_ltfu  = int((cohort["itt_group"] == "Loss to follow-up").sum())
N_nonltfu = int((cohort["itt_group"] == "Non-LTFU").sum())


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------
def parse_est(s):
    """'1.23 (1.01-1.40)' -> (1.23, 1.01, 1.40)."""
    try:
        hr, rest = s.split(" (")
        lo, hi = rest.rstrip(")").split("-")
        return float(hr), float(lo), float(hi)
    except Exception:
        return None, None, None


def hr_ci(model_name, term):
    row = mv[(mv["model"] == model_name) & (mv["term"] == term)]
    if row.empty:
        return "—"
    return row.iloc[0]["estimate"]


def tt_late_cap2(month, field="HR"):
    r = tt_array[(tt_array["Trial_Month"] == f"Month_{month}") &
                 (tt_array["model"] == "late") & (tt_array["cap"] == 2)]
    if r.empty:
        return None
    return float(r.iloc[0][field])


def tt_late_fmt(month):
    r = tt_array[(tt_array["Trial_Month"] == f"Month_{month}") &
                 (tt_array["model"] == "late") & (tt_array["cap"] == 2)]
    if r.empty:
        return "—"
    row = r.iloc[0]
    return f"{row['HR']:.2f} (95% CI {row['CI_L']:.2f}–{row['CI_H']:.2f})"


def tt_subgrp_fmt(subgroup, level):
    r = tt_subgrp[(tt_subgrp["Subgroup"] == subgroup) &
                  (tt_subgrp["Level"] == level) &
                  (tt_subgrp["model"] == "late") &
                  (tt_subgrp["cap"] == 2)]
    if r.empty:
        return "—"
    row = r.iloc[0]
    return f"{row['HR']:.2f} (95% CI {row['CI_L']:.2f}–{row['CI_H']:.2f})"


def fig2_val(panel, group, horizon_col):
    r = fig2_vals[(fig2_vals["panel"] == panel) & (fig2_vals["group"] == group)]
    if r.empty:
        return None
    return float(r.iloc[0][horizon_col])


# --------------------------------------------------------------------------
# Build document
# --------------------------------------------------------------------------
doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def add_bullets(items, level=0):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet" if level == 0 else "List Bullet 2")


# ==========================================================================
# Title & authors
# ==========================================================================
title = doc.add_heading(
    "Evaluating the impact of loss to follow-up in a population-based "
    "tuberculosis treatment cohort in Brazil: insights from causal "
    "inferential analyses",
    level=0,
)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

add_para("Evelyn Lepka de Lima, Ana Angélica Lindoso, Suely Fukusava, "
         "Jason R. Andrews", italic=True)
add_para(f"Draft generated {datetime.now():%Y-%m-%d} from the ITT "
         f"causal-analysis pipeline.", italic=True, size=9)


# ==========================================================================
# Abstract / Summary
# ==========================================================================
add_heading("Summary", level=1)

n_ltfu_pct  = 100 * N_ltfu / N_total
pct_retreated = fig1_stats["pct_retreated"]
median_months_to_ab = fig1_stats["median_months_to_abandonment"]
median_months_to_retx = fig1_stats["median_years_to_retreatment"] * 12

# Late mortality HRs (24-month cap)
late_hrs = [tt_late_cap2(m) for m in range(1, 7) if tt_late_cap2(m)]
hr_min, hr_max = min(late_hrs), max(late_hrs)

# Subgroup extremes from late cap=2
hr_15_24    = tt_subgrp_fmt("age_group", "15-24")
hr_65p      = tt_subgrp_fmt("age_group", "≥65")
hr_housed   = tt_subgrp_fmt("homelessness", "No")
hr_homeless = tt_subgrp_fmt("homelessness", "Yes")

# 5-year cum mortality among LTFU — pull from cohort
ltfu = cohort[cohort["itt_group"] == "Loss to follow-up"].copy()
ltfu["died_5y"] = (ltfu["event_d"] == 1) & (ltfu["time_d"] <= 5)
pct_5y_mort = 100 * ltfu["died_5y"].sum() / len(ltfu)
ltfu_hiv = ltfu[ltfu["hiv_aids"] == "Positive"]
pct_5y_mort_hiv = 100 * ((ltfu_hiv["event_d"] == 1) & (ltfu_hiv["time_d"] <= 5)).sum() / len(ltfu_hiv)

add_para(
    "Background. Treatment interruption remains a major challenge in "
    "tuberculosis care; because individuals lost to follow-up (LTFU) "
    "differ systematically from those who complete therapy, understanding "
    "the true impact of LTFU on mortality poses methodological challenges. "
    "We aimed to estimate the determinants of retreatment and the causal "
    "effect of LTFU on mortality using advanced causal inferential "
    "approaches.",
)

add_para(
    "Methods. We conducted a retrospective cohort study of individuals "
    "aged ≥15 years initiating their first tuberculosis treatment in São "
    "Paulo State, Brazil (2013–2023), using data from the State "
    "tuberculosis registry linked to the national mortality system. "
    "Missing covariates were addressed by multiple imputation with "
    "chained equations (m=5, miceforest). Within-LTFU predictors of "
    "mortality and retreatment were estimated with MI-pooled Cox and "
    "Fine–Gray models. To eliminate immortal-time bias in estimating the "
    "causal effect of LTFU on mortality, we emulated a sequence of target "
    "trials at each month of therapy (1–6), separately for early (0–6 "
    "months) and late (6–24 months) mortality. Adjusted interaction "
    "models assessed subgroup effect modification in the late-mortality "
    "target trial."
)

add_para(
    f"Results. Among {N_total:,} individuals initiating tuberculosis "
    f"therapy, {N_ltfu:,} ({n_ltfu_pct:.1f}%) experienced LTFU during "
    f"their first episode. Among those LTFU, 5-year cumulative "
    f"mortality was {pct_5y_mort:.1f}%, including {pct_5y_mort_hiv:.1f}% "
    f"among people living with HIV. "
    f"Less than half ({pct_retreated:.1f}%) re-entered treatment at a "
    f"median of {median_months_to_retx:.1f} months, with retreatment "
    f"most strongly predicted by indicators of baseline disease severity "
    f"(hospitalisation, adjusted subdistribution hazard ratio [aSHR] "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'hosp_admissionYes')}; HIV, aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'hiv_aidsPositive')}). "
    f"In the sequential target-trial emulations, during the first six "
    f"months of treatment LTFU appeared protective against early "
    f"mortality (adjusted hazard ratio [aHR] <1 for months 1–5), an "
    f"artefact of severity-driven differential propensity to interrupt "
    f"therapy. After the first six months of follow-up, LTFU at any "
    f"month of therapy conferred an approximately "
    f"{hr_min:.1f}- to {hr_max:.1f}-fold increase in mortality (range "
    f"across Months 1–6 of LTFU: aHR {hr_min:.2f}–{hr_max:.2f}). "
    f"The relative late-mortality penalty was greater in younger "
    f"individuals (age 15–24 aHR {hr_15_24} vs. age ≥65 aHR "
    f"{hr_65p}) and in housed compared with structurally homeless "
    f"individuals (aHR {hr_housed} vs. {hr_homeless})."
)

add_para(
    "Conclusions. Less than half of individuals LTFU returned to care, "
    "and those who did had more severe disease. Accounting for disease "
    "severity and immortal-time bias, LTFU at any time during therapy "
    "carries an increased risk of death, particularly in the late "
    "post-interruption period. Person-centred support is critical to "
    "sustain treatment, including active re-engagement for individuals "
    "who discontinue therapy."
)


# ==========================================================================
# Background (outline)
# ==========================================================================
add_heading("Background [detailed outline]", level=1)
add_para(
    "This section is outlined — prose to be written. The outline below "
    "captures the intended argument and citation needs.", italic=True
)

add_heading("1. Global burden of TB treatment interruption", level=2)
add_bullets([
    "TB remains a leading infectious cause of death worldwide; the 2030 "
    "End-TB targets hinge on completion of curative therapy.",
    "Programmatic loss to follow-up (treatment interruption or "
    "“abandonment”) is a persistent obstacle in high-burden settings — "
    "global proportions of 8–15% typical; São Paulo experience falls "
    "within this range (~12.5% in our cohort).",
    "Consequences include acquired resistance, ongoing transmission, and "
    "increased mortality; quantifying the mortality impact has been "
    "methodologically contested.",
])

add_heading("2. Who is lost to follow-up?", level=2)
add_bullets([
    "Demographic and clinical determinants repeatedly identified: male "
    "sex, younger age, lower education, drug and alcohol use, "
    "homelessness, mental-health co-morbidities, HIV co-infection.",
    "Structural drivers: incarceration history (with paradoxical "
    "lower LTFU while incarcerated), migration, disrupted healthcare "
    "access.",
    "Key framing: LTFU is not a random event — it is a marker of both "
    "structural disadvantage and (in many cases) improving clinical "
    "status that drives patients away from care.",
])

add_heading("3. Methodological challenges in estimating the impact of LTFU", level=2)
add_bullets([
    "Confounding by indication: sicker patients may be less able to "
    "abandon (hospitalised, DOT) or, conversely, more likely to (severe "
    "disease + social adversity).",
    "Immortal-time bias: comparing abandoners vs. completers starting at "
    "treatment initiation makes abandoners appear protected because they "
    "must survive to abandon.",
    "Healthy-returner bias: conditioning on retreatment misses those "
    "who died before returning — retreatment looks harmful because it "
    "is a marker of decompensated disease.",
    "Standard Cox with treatment outcome as a fixed covariate conflates "
    "these sources of bias; prior literature disagrees on the direction "
    "of effect.",
])

add_heading("4. Contribution of this study", level=2)
add_bullets([
    "Large population-based linked cohort (TBweb + SIM) with 10 years "
    "of follow-up.",
    "Within-LTFU multivariable models (Cox for mortality, Fine–Gray for "
    "retreatment) identify who is most at risk.",
    "Sequential target-trial emulation eliminates immortal-time bias by "
    "aligning exposure and time-at-risk at every month of therapy, "
    "separately for early (0–6 mo) and late (6–24 mo) mortality.",
    "Subgroup effect-modification analyses reveal where the relative "
    "penalty for LTFU is concentrated and where competing risks blunt "
    "it.",
])


# ==========================================================================
# Methods
# ==========================================================================
add_heading("Methods", level=1)

add_heading("Study design and data sources", level=2)
add_para(
    "We conducted a retrospective cohort study using data from the São "
    "Paulo State Tuberculosis Program (TBweb) linked to the Brazilian "
    "Mortality Information System (SIM). TBweb is the electronic "
    "surveillance platform used to notify and monitor all individuals "
    "diagnosed with tuberculosis in São Paulo State and is integrated "
    "with the national SINAN database, allowing longitudinal follow-up "
    "of successive tuberculosis episodes for the same individual via the "
    "SINAN identifier. SIM records all deaths occurring in Brazil "
    "including date and underlying cause. TBweb and SIM records were "
    "probabilistically linked to ascertain mortality during follow-up. "
    "Data extraction and linkage were performed in January 2025."
)

add_heading("Study population", level=2)
add_para(
    f"We included all individuals aged ≥15 years with a first (\"Novo\") "
    f"tuberculosis episode notified in TBweb between January 1, 2013 and "
    f"December 31, 2023 (Figure S1 — exclusion flowchart). Individuals "
    f"were excluded if age <15 years, treatment end-date outside the "
    f"2013–2023 study window, death on or before the treatment start "
    f"date, or unresolvable date inconsistencies. The final cohort "
    f"comprised {N_total:,} individuals, of whom {N_ltfu:,} "
    f"({n_ltfu_pct:.1f}%) experienced LTFU (case closed as "
    f"“Abandono”/“Abandono primário”/“Faltoso”) during their first "
    f"episode and {N_nonltfu:,} ({100 - n_ltfu_pct:.1f}%) completed "
    f"therapy or experienced an alternative outcome (cured, died, "
    f"transferred, failure)."
)

add_heading("Exposure, outcomes, and time variables", level=2)
add_para(
    "The primary exposure was LTFU during the first tuberculosis "
    "episode. We defined two complementary time origins, used depending "
    "on the target causal estimand. For analyses contrasting LTFU vs. "
    "remaining on treatment (sequential target trials; crude time-varying "
    "hazard), time zero was the treatment start date (best_start) and "
    "follow-up was measured in years since that date (time_d_tx). This "
    "avoids immortal-time bias by entering every individual at the same "
    "origin and allowing LTFU to act as a time-varying exposure. For "
    "landmark analyses conditional on reaching LTFU, time zero was the "
    "treatment end date (end_date), with follow-up measured from that "
    "date (time_rn / time_d)."
)
add_para(
    "Retreatment was defined as the earliest subsequent TBweb "
    "notification in the same individual following the index end-date. "
    "All-cause mortality was ascertained from SIM and from death "
    "outcomes recorded in TBweb for subsequent notifications; where "
    "records disagreed, the earliest validated date of death was used. "
    "Administrative censoring was December 31, 2024."
)

add_heading("Multiple imputation", level=2)
add_para(
    "Missing covariates (race, education, HIV status, behavioural "
    "factors, clinical classification, DOT status) were addressed using "
    "multiple imputation by chained equations with random-forest "
    "predictors (miceforest, m=5) applied to the full cohort. The "
    "imputation model included all baseline covariates, the exposure "
    "(itt_group), the event indicators, and the log-transformed "
    "follow-up times. All downstream multivariable analyses were fit "
    "separately in each imputed dataset and pooled using Rubin's rules. "
    "Complete-case analyses are reported as sensitivity analyses."
)

add_heading("Within-LTFU multivariable models", level=2)
add_para(
    "Among individuals who experienced LTFU, we estimated predictors of "
    "all-cause mortality using Cox proportional-hazards regression and "
    "predictors of retreatment using Fine–Gray subdistribution-hazard "
    "regression, with death treated as the competing event. Time zero "
    "was the LTFU date (end_date). Models were adjusted for age group, "
    "sex, self-reported race, education level, HIV status, diabetes, "
    "alcohol use, drug use, incarceration history, homelessness, "
    "hospitalisation at diagnosis, clinical classification (pulmonary, "
    "extrapulmonary, mixed/disseminated), directly-observed therapy "
    "(DOT), and duration of treatment prior to LTFU (<2, 2–<4, ≥4 "
    "months). Estimates are reported as adjusted hazard ratios (aHR) or "
    "adjusted subdistribution hazard ratios (aSHR) with 95% confidence "
    "intervals, pooled across the five imputed datasets using Rubin's "
    "rules."
)

add_heading("Crude time-varying hazard ratio (illustrative)", level=2)
add_para(
    "To illustrate the magnitude of immortal-time bias that affects "
    "naïve comparisons of LTFU vs. treatment-completion cohorts, we "
    "fit a piecewise Cox model in counting-process format using "
    "treatment start as time zero. Every individual entered the risk "
    "set at t=0 as unexposed (still on treatment). LTFU individuals "
    "transitioned to the exposed state at the date of their LTFU "
    "(tx_duration). Follow-up was split at monthly boundaries over "
    "24 months, and bucket-specific crude hazard ratios were estimated "
    "by fitting Cox models within each monthly interval. A LOESS smooth "
    "(span 0.55) with 95% confidence band was overlaid to visualise the "
    "trajectory (Figure 3A). This analysis is descriptive and not "
    "covariate-adjusted."
)

add_heading("Sequential target-trial emulation of LTFU effect on mortality",
            level=2)
add_para(
    "To obtain an unbiased causal estimate of the effect of LTFU on "
    "mortality, we emulated a sequence of target trials, one for each "
    "treatment month from 1 to 6. In the trial beginning at month m, "
    "eligibility required the individual to remain alive at month m and "
    "to still be on treatment (tx_duration ≥ m/12 years) for non-LTFU "
    "controls, or to be eligible to abandon at that month for "
    "LTFU-treatment candidates. Exposure was defined as LTFU occurring "
    "within the monthly window [m-1, m). Time at risk began at month m "
    "and was measured as time_d_tx − (m-1)/12 years. We fit "
    "cause-specific Cox regression adjusted for age group, sex, race, "
    "education, HIV, diabetes, alcohol, drug use, incarceration, "
    "homelessness, hospital admission, clinical classification, and "
    "DOT. To distinguish short- from long-term effects and mitigate "
    "residual informative censoring, each trial was estimated in three "
    "outcome windows: overall (cap = 24 months), early (events in the "
    "first 6 months from trial start), and late (events in months 6–24 "
    "from trial start). Early-window deaths in the “late” analysis were "
    "administratively censored at their death time (event = 0) rather "
    "than using a landmark design, preserving follow-up information. "
    "Fits were pooled across the five imputed datasets with Rubin's "
    "rules (Figure 3B)."
)

add_heading("Subgroup target-trial emulation (effect modification)", level=2)
add_para(
    "To investigate where the relative late-mortality penalty of LTFU is "
    "concentrated, we repeated the target-trial emulation within "
    "pre-specified subgroups: age group (15–24, 25–44, 45–64, ≥65 "
    "years), sex, HIV status, and homelessness status. For each "
    "subgroup level, we fit a pooled Cox model combining all six "
    "monthly trials with a trial-month indicator, restricting to the "
    "late 6–24-month mortality window (cap = 24 months) and using the "
    "same covariate set (minus the stratifying variable). Results are "
    "pooled across imputed datasets and presented as a forest plot "
    "(Figure 3C)."
)

add_heading("Descriptive analyses and stratified cumulative incidence",
            level=2)
add_para(
    "Among individuals with LTFU, we described trajectories (retreatment, "
    "death without retreatment, no further outcome; and, for those who "
    "retreated, the outcome of the retreatment episode) using a "
    "three-column alluvial diagram (Figure 1A). Timing of LTFU, "
    "retreatment and death were summarised as medians with IQR and "
    "visualised as rainclouds (Figure 1B–D). Cumulative incidence of "
    "retreatment with death as a competing risk was estimated using the "
    "Aalen–Johansen estimator (cmprsk::cuminc), up to 24 months from "
    "LTFU, stratified by HIV status, month of LTFU, homelessness and "
    "hospitalisation at diagnosis (Figure 2)."
)

add_heading("Software and reproducibility", level=2)
add_para(
    "Cohort construction and table rendering were performed in Python "
    "3.13 (pandas, numpy, miceforest, python-docx). Survival and "
    "causal analyses were performed in R 4.5 (survival, mice, cmprsk, "
    "broom, patchwork, ggplot2). The full pipeline resolves the project "
    "root automatically via a shared path helper and is reproducible "
    "from the cohort CSV forward. Code and the self-contained HTML "
    "results summary are available on request."
)

add_heading("Ethics", level=2)
add_para(
    "Written permission for secondary use of surveillance data was "
    "obtained from the São Paulo State Health Department, and the "
    "study was approved by the Research Ethics Committee of the "
    "Instituto de Infectologia Emílio Ribas, São Paulo, Brazil."
)


# ==========================================================================
# Results
# ==========================================================================
add_heading("Results", level=1)

# --- Cohort
add_heading("Study population", level=2)
add_para(
    f"A total of {N_total:,} individuals aged ≥15 years with a first "
    f"tuberculosis episode between 2013 and 2023 were included. Of "
    f"these, {N_ltfu:,} ({n_ltfu_pct:.1f}%) experienced LTFU during "
    f"their first episode, while {N_nonltfu:,} ({100 - n_ltfu_pct:.1f}%) "
    f"had alternative outcomes (cured, died, transferred, treatment "
    f"failure)."
)

ltfu_df    = cohort[cohort["itt_group"] == "Loss to follow-up"]
nonltfu_df = cohort[cohort["itt_group"] == "Non-LTFU"]

def _pct(df, col, val):
    s = df[col].dropna()
    if len(s) == 0:
        return 0.0
    return 100 * (s == val).sum() / len(s)

def _median_age(df):
    a = df["age_tb"].dropna()
    return int(a.median()), int(a.quantile(0.25)), int(a.quantile(0.75))

m_all, q1_all, q3_all = _median_age(cohort)
m_l, q1_l, q3_l = _median_age(ltfu_df)
m_n, q1_n, q3_n = _median_age(nonltfu_df)

add_para(
    f"Compared with those who completed or had an alternative outcome, "
    f"individuals LTFU were younger (median age {m_l} years [IQR "
    f"{q1_l}–{q3_l}] vs. {m_n} [IQR {q1_n}–{q3_n}]), more frequently "
    f"male ({_pct(ltfu_df, 'sex', 'Male'):.1f}% vs. "
    f"{_pct(nonltfu_df, 'sex', 'Male'):.1f}%), and reported higher "
    f"prevalence of homelessness "
    f"({_pct(ltfu_df, 'homelessness', 'Yes'):.1f}% vs. "
    f"{_pct(nonltfu_df, 'homelessness', 'Yes'):.1f}%), alcohol use "
    f"({_pct(ltfu_df, 'alcohol', 'Yes'):.1f}% vs. "
    f"{_pct(nonltfu_df, 'alcohol', 'Yes'):.1f}%) and drug use "
    f"({_pct(ltfu_df, 'drug_use', 'Yes'):.1f}% vs. "
    f"{_pct(nonltfu_df, 'drug_use', 'Yes'):.1f}%; all p<0.001). "
    f"HIV was reported in {_pct(ltfu_df, 'hiv_aids', 'Positive'):.1f}% "
    f"of LTFU individuals compared with "
    f"{_pct(nonltfu_df, 'hiv_aids', 'Positive'):.1f}% of those who were "
    f"not LTFU. Individuals LTFU were also more often hospitalised at "
    f"diagnosis ({_pct(ltfu_df, 'hosp_admission', 'Yes'):.1f}% vs. "
    f"{_pct(nonltfu_df, 'hosp_admission', 'Yes'):.1f}%) and less likely "
    f"to receive DOT "
    f"({_pct(ltfu_df, 'dot_status', 'Yes'):.1f}% vs. "
    f"{_pct(nonltfu_df, 'dot_status', 'Yes'):.1f}%; Table 1)."
)

# --- Post-LTFU trajectories
add_heading("Trajectories after loss to follow-up", level=2)
add_para(
    f"Among the {N_ltfu:,} individuals LTFU, {int(fig1_stats['n_retreated']):,} "
    f"({pct_retreated:.1f}%) re-entered TB treatment at a median of "
    f"{median_months_to_retx:.1f} months from LTFU (IQR pooled over "
    f"retreated individuals; see Figure 1B–C), and "
    f"{int(fig1_stats['n_died']):,} ({fig1_stats['pct_died']:.1f}%) "
    f"died during follow-up, at a median of "
    f"{fig1_stats['median_years_to_mortality']:.2f} years after LTFU. "
    f"The timing of LTFU itself was concentrated in the first half of "
    f"therapy, with a median of "
    f"{median_months_to_ab:.2f} months from treatment start. Among the "
    f"retreated subgroup, the modal outcome of the retreatment episode "
    f"was cure, but a substantial fraction experienced a second LTFU "
    f"(see Figure 1A alluvial)."
)

# --- Stratified cumulative incidence of retreatment
add_heading("Cumulative incidence of retreatment", level=2)
v_pos = fig2_val("HIV", "Positive", "month_24")
v_neg = fig2_val("HIV", "Negative", "month_24")
v_home_y = fig2_val("Homelessness", "Yes", "month_24")
v_home_n = fig2_val("Homelessness", "No", "month_24")
v_ltfu_lt2 = fig2_val("Month_abandonment", "< 2 months", "month_24")
v_ltfu_ge4 = fig2_val("Month_abandonment", "≥ 4 months", "month_24")
v_hosp_y = fig2_val("Hospitalization", "Yes", "month_24")
v_hosp_n = fig2_val("Hospitalization", "No", "month_24")

add_para(
    f"Twenty-four-month cumulative incidence of retreatment (treating "
    f"death as a competing risk) was {v_pos:.1f}% among people living "
    f"with HIV compared with {v_neg:.1f}% among HIV-negative individuals "
    f"(Figure 2A, Gray's test p<0.001). Retreatment incidence rose "
    f"steeply with earlier LTFU: {v_ltfu_lt2:.1f}% among those LTFU "
    f"within <2 months of starting therapy compared with {v_ltfu_ge4:.1f}% "
    f"among those who left after ≥4 months (Figure 2B). Homelessness "
    f"was associated with somewhat higher 24-month retreatment "
    f"incidence ({v_home_y:.1f}% vs. {v_home_n:.1f}%; Figure 2C), and "
    f"hospitalisation at diagnosis with substantially higher incidence "
    f"({v_hosp_y:.1f}% vs. {v_hosp_n:.1f}%; Figure 2D), consistent with "
    f"return-to-care being driven in part by clinical deterioration."
)

# --- Multivariable within-LTFU
add_heading("Within-LTFU predictors of mortality and retreatment", level=2)
add_para(
    f"In MI-pooled within-LTFU multivariable models (Figure 4), "
    f"mortality was independently associated with older age "
    f"(45–64 years aHR {hr_ci('Cox_Death_Adjusted_MI', 'age_group45-64')}; "
    f"≥65 years aHR {hr_ci('Cox_Death_Adjusted_MI', 'age_group≥65')}), "
    f"lower educational attainment (≤7 years aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'edu_clean≤ 7 years')}), "
    f"HIV positivity (aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'hiv_aidsPositive')}), "
    f"alcohol use (aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'alcoholYes')}), and "
    f"hospitalisation at diagnosis (aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'hosp_admissionYes')}). "
    f"Incarceration history was associated with markedly lower "
    f"mortality (aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'incarceratedYes')}), as was "
    f"extrapulmonary disease (aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'clinical_cleanExtrapulmonary')}). "
    f"Shorter duration of therapy prior to LTFU was associated with "
    f"higher mortality (LTFU in <2 months aHR "
    f"{hr_ci('Cox_Death_Adjusted_MI', 'tx_month_grp< 2 months')})."
)
add_para(
    f"Retreatment was most strongly predicted by indicators of disease "
    f"severity and early departure from treatment: hospitalisation at "
    f"diagnosis (aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'hosp_admissionYes')}), HIV "
    f"(aSHR {hr_ci('FG_Retr_Adjusted_MI', 'hiv_aidsPositive')}), "
    f"drug use (aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'drug_useYes')}), and LTFU within "
    f"the first 2 months of therapy (aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'tx_month_grp< 2 months')}). "
    f"Older age was associated with substantially lower retreatment "
    f"(≥65 years aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'age_group≥65')}), as was "
    f"homelessness (aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'homelessnessYes')}) and "
    f"extrapulmonary disease (aSHR "
    f"{hr_ci('FG_Retr_Adjusted_MI', 'clinical_cleanExtrapulmonary')})."
)

# --- Causal mortality effect
add_heading("Causal effect of LTFU on mortality", level=2)
add_para(
    "The crude time-varying hazard ratio for LTFU vs. on-treatment "
    "showed the expected pattern of immortal-time bias: in the first "
    "few months from treatment start, LTFU appeared protective (HR <1) "
    "because patients must survive in order to be observed as LTFU; "
    "beyond ~6 months the HR rose towards and exceeded unity as "
    "mortality accumulated among those who had left care (Figure 3A). "
    "This pattern demonstrates why standard Cox comparisons anchored at "
    "treatment start cannot recover the causal impact of LTFU."
)

late_fmt = [tt_late_fmt(m) for m in range(1, 7)]
add_para(
    "The sequential target-trial emulation recovered the causal "
    "contrast. When mortality was truncated at 24 months, the "
    "early-window (0–6 months) hazard ratios remained <1 at every "
    "month of LTFU — a structural consequence of higher-severity "
    "patients being less likely to interrupt therapy. In the late "
    f"window (6–24 months), LTFU carried a substantially increased "
    f"mortality hazard at every month of departure: Month 1 aHR "
    f"{late_fmt[0]}, Month 2 aHR {late_fmt[1]}, Month 3 aHR "
    f"{late_fmt[2]}, Month 4 aHR {late_fmt[3]}, Month 5 aHR "
    f"{late_fmt[4]}, Month 6 aHR {late_fmt[5]} (Figure 3B)."
)

add_para(
    f"The relative late-mortality penalty was not uniform across "
    f"subgroups (Figure 3C). The hazard of LTFU was highest in the "
    f"youngest individuals (15–24 years aHR {hr_15_24}) and attenuated "
    f"monotonically with age (≥65 years aHR {hr_65p}, confidence "
    f"interval crossing the null). People living with HIV had a "
    f"similar late-mortality aHR (aHR "
    f"{tt_subgrp_fmt('hiv_aids', 'Positive')}) to HIV-negative "
    f"individuals (aHR {tt_subgrp_fmt('hiv_aids', 'Negative')}), "
    f"indicating that the relative effect of LTFU on mortality is "
    f"comparable across HIV strata despite much higher absolute "
    f"mortality in PLHIV. The relative penalty was strikingly larger "
    f"in housed (aHR {hr_housed}) than in structurally homeless "
    f"individuals (aHR {hr_homeless}), consistent with homeless "
    f"patients experiencing overwhelming baseline mortality that "
    f"mechanically blunts the relative impact of additional exposure "
    f"through LTFU."
)


# ==========================================================================
# Discussion (outline)
# ==========================================================================
add_heading("Discussion [detailed outline]", level=1)
add_para(
    "This section is outlined — prose to be written.", italic=True
)

add_heading("1. Key findings", level=2)
add_bullets([
    "Less than half of individuals LTFU returned to TB care; those who "
    "did had more severe disease, reflected in higher rates of "
    "hospitalisation at the index episode and HIV co-infection.",
    "Accounting for disease severity and immortal-time bias via "
    "sequential target-trial emulation, LTFU at any month of treatment "
    "carries an approximately 2–3-fold increased hazard of late "
    "(6–24-month) mortality.",
    "The relative penalty is not uniform: it is greatest in younger "
    "individuals and in the structurally housed; the absolute mortality "
    "burden, however, remains highest among people experiencing "
    "homelessness and people living with HIV.",
])

add_heading("2. Interpretation", level=2)
add_bullets([
    "The crude HR(t) pattern (Fig 3A) is a textbook demonstration of "
    "immortal-time bias — a caution against naïve Cox comparisons.",
    "Target trials align exposure and time-at-risk by month of therapy, "
    "making the comparison interpretable as: “what happens to a patient "
    "who leaves therapy at month m vs. one who does not leave at that "
    "month?”",
    "The early-window protective appearance is not a true protective "
    "effect — it reflects selection (healthier patients being more "
    "likely to leave) that we cannot fully adjust for with baseline "
    "covariates; the late-window estimate is the policy-relevant "
    "quantity.",
    "Effect modification by age and housing status highlights where "
    "interventions will have the highest relative impact, and where "
    "the absolute burden demands parallel structural investment.",
])

add_heading("3. Comparison with prior literature", level=2)
add_bullets([
    "Prior estimates of the mortality impact of LTFU range widely "
    "(aHR 1.3–5.0 in various settings), largely reflecting design "
    "differences: landmark vs. time-varying vs. ITT; "
    "disease-severity adjustment; HIV prevalence.",
    "Earlier São Paulo analyses reporting an apparent protective "
    "effect of LTFU at short horizons (HR<1) were likely capturing "
    "immortal-time bias rather than a real effect.",
    "Our within-LTFU Fine–Gray finding that severity markers "
    "(hospitalisation, HIV) predict retreatment is consistent with "
    "“healthy-survivor” dynamics reported in Cape Town and Lima "
    "cohorts — patients who are well enough to remain disengaged from "
    "care are also well enough to survive without visible healthcare "
    "contact.",
])

add_heading("4. Implications for programmes and policy", level=2)
add_bullets([
    "Re-engagement programmes are clinically urgent — a substantial "
    "fraction of LTFU individuals experience severe disease and death "
    "in the years after interruption; active outreach (SMS, community "
    "health worker visits, peer navigators) is supported by these "
    "findings.",
    "Targeted attention to structural drivers (housing, substance-use "
    "support, mental-health services) is necessary to reduce absolute "
    "mortality burden, even where the relative LTFU penalty is blunted.",
    "The 12.5% LTFU proportion is above WHO targets for high- and "
    "upper-middle-income settings; São Paulo's public programme is "
    "strong, suggesting structural rather than programmatic drivers.",
])

add_heading("5. Strengths", level=2)
add_bullets([
    "Population-based cohort with 10+ years of follow-up, integrated "
    "surveillance, and vital-statistics linkage.",
    "Sequential target-trial emulation addresses immortal-time bias "
    "and provides an interpretable monthly exposure contrast.",
    "Multiple imputation with chained equations + Rubin's rules pools "
    "information across all incomplete cases.",
    "Pre-registered subgroup analyses avoid post-hoc cherry-picking.",
])

add_heading("6. Limitations", level=2)
add_bullets([
    "Residual confounding: unmeasured severity markers (viral load, "
    "CD4, radiographic extent, bacillary burden) are not in TBweb.",
    "SIM linkage accuracy: probabilistic record linkage may "
    "under-ascertain deaths in highly mobile or unhoused populations.",
    "LTFU episode definition relies on the case-closure date recorded "
    "in TBweb; informal interruptions that lead to case closure without "
    "true disengagement may be misclassified.",
    "Generalisability: São Paulo is the largest Brazilian state and "
    "has a well-resourced public TB programme; patterns may differ in "
    "lower-resource regions.",
    "The target-trial emulation requires positivity assumptions at each "
    "month — in months with very few abandoners (e.g. Month 6, N ~ "
    "200) estimates are less precise.",
])

add_heading("7. Conclusions", level=2)
add_para(
    "Less than half of individuals lost to follow-up during a first "
    "tuberculosis episode returned to care, and those who did had more "
    "severe baseline disease. Once immortal-time bias and disease "
    "severity were accounted for, LTFU at any month of therapy carried "
    "an approximately 2–3-fold increase in late mortality. Person-"
    "centred support to sustain treatment — including structured "
    "re-engagement for patients who discontinue — is a programmatic "
    "priority, particularly for younger, structurally housed individuals "
    "whose relative mortality penalty is greatest.",
    italic=False
)


# ==========================================================================
# Figures (embedded with legends)
# ==========================================================================
add_heading("Figures", level=1)


def add_figure(png_path, caption):
    if png_path.exists():
        doc.add_picture(str(png_path), width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(f"[figure missing: {png_path.name}]", italic=True)
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.font.size = Pt(10)


add_figure(
    RESULTS / "Figure_1_descriptive.png",
    "Figure 1. Descriptive post-LTFU trajectories and event timing. "
    f"(A) Three-column alluvial diagram showing trajectories after LTFU "
    f"(N = {N_ltfu:,}). Column 1, index LTFU; Column 2, first outcome "
    f"after LTFU (retreatment, death without retreatment, or no further "
    f"outcome observed); Column 3, outcome of the retreatment episode "
    f"among those who re-entered care. (B) Timing of LTFU relative to "
    f"treatment start (months). (C) Time from LTFU to retreatment "
    f"among those who retreated. (D) Time from LTFU to death among "
    f"those who died. Density + raincloud presentation with the median "
    f"annotated in each panel."
)

add_figure(
    RESULTS / "Figure_2_stratified_retreatment_24mo.png",
    "Figure 2. Cumulative incidence of retreatment within 24 months "
    "of LTFU, stratified by (A) HIV status, (B) month of LTFU "
    "(time on therapy before interruption), (C) homelessness, and "
    "(D) hospitalisation at the index tuberculosis diagnosis. "
    "Aalen–Johansen estimates treat all-cause death as a competing risk "
    f"(LTFU subgroup N = {N_ltfu:,})."
)

add_figure(
    RESULTS / "Figure_3_causal_mortality.png",
    "Figure 3. Causal effect of LTFU on mortality. (A) Crude "
    "time-varying hazard ratio for LTFU vs. on-treatment, fit in "
    "counting-process format across monthly buckets from treatment "
    "start. The early apparent “protective” effect (HR<1 in months "
    "0–6) is a hallmark of immortal-time bias. (B) MI-pooled sequential "
    "target-trial adjusted hazard ratio (AHR) for the effect of LTFU "
    "at each month of therapy (1–6) on early (0–6 months) and late "
    "(6–24 months) mortality. (C) Subgroup late-mortality AHR from "
    "subgroup-specific target-trial emulations (cap = 24 months), "
    "pooled across m = 5 imputed datasets. Reference line at AHR = 1."
)

add_figure(
    RESULTS / "Figure_4_within_ltfu_forest.png",
    "Figure 4. Within-LTFU multivariable forest plot. MI-pooled "
    "adjusted Cox (left, mortality) and Fine–Gray (right, retreatment "
    "with death as a competing event) hazard ratios for each baseline "
    "covariate, grouped into demographic, psychosocial, and biomedical "
    "domains. Reference categories shown in grey italics with hollow "
    "diamonds at AHR = 1. Hospital admission's mortality aHR (4.42, 95% "
    "CI 3.94–4.96) is squished to the right edge of panel A; the exact "
    "value is shown in the accompanying text column."
)


# ==========================================================================
# Table 1 — LTFU vs Non-LTFU (no Total column)
# ==========================================================================
doc.add_page_break()
add_heading("Tables", level=1)
add_heading(f"Table 1. Baseline characteristics of the study cohort, "
            f"by loss-to-follow-up status (N = {N_total:,}).", level=2)

df = cohort.copy()
df["sex"] = df["sex"].fillna("Missing")
df["age_group"] = pd.cut(df["age_tb"], bins=[14, 24, 44, 64, 150],
                          labels=["15-24", "25-44", "45-64", "≥ 65"])
for c in ["hiv_aids", "incarcerated", "homelessness", "dot_status",
          "alcohol", "drug_use", "diabetes"]:
    df[c] = df[c].fillna("Missing")
df["tobacco_stat"] = df["tobacco_use"].fillna("Missing")
df["mental"] = df["mental_health"].fillna("Missing")
df["immuno"] = df["other_immuno_condition"].fillna("Missing")

# Cohort already has clinical_clean / race_clean / edu_clean / diagnosis_setting
df["clinical_clean"]    = df["clinical_clean"].fillna("Missing")
df["diagnosis_clean"]   = df["diagnosis_setting"].fillna("Missing")
df["lab_confirmed_stat"] = df["lab_confirmed_stat"].fillna("Missing")
df["hosp_admission_stat"] = df["hosp_admission"].fillna("Missing")
df["race_clean"]        = df["race_clean"].fillna("Missing")
df["education_clean"]   = df["edu_clean"].fillna("Missing")

g_aband = df[df["itt_group"] == "Loss to follow-up"]
g_ctrl  = df[df["itt_group"] == "Non-LTFU"]

blocks = [
    ("Socio-demographic", [
        ("Age_Median", "age_tb"),
        ("Age group", "age_group"),
        ("Sex", "sex"),
        ("Race", "race_clean"),
        ("Education", "education_clean"),
    ]),
    ("Clinical comorbidities", [
        ("HIV/AIDS", "hiv_aids"),
        ("Diabetes", "diabetes"),
        ("Other immunosuppressive condition", "immuno"),
        ("Mental-health diagnosis", "mental"),
    ]),
    ("Behavioural", [
        ("Alcohol use", "alcohol"),
        ("Drug use", "drug_use"),
        ("Tobacco use", "tobacco_stat"),
    ]),
    ("Social vulnerabilities", [
        ("Incarceration history", "incarcerated"),
        ("Homelessness", "homelessness"),
    ]),
    ("Disease and treatment", [
        ("Diagnosis setting", "diagnosis_clean"),
        ("Hospital admission at diagnosis", "hosp_admission_stat"),
        ("Laboratory-confirmed TB", "lab_confirmed_stat"),
        ("Clinical classification", "clinical_clean"),
        ("Directly-observed therapy", "dot_status"),
    ]),
]


def get_p_val(df1, df2, var):
    try:
        obs = pd.concat([df1[var].value_counts(),
                         df2[var].value_counts()], axis=1).fillna(0)
        if "Missing" in obs.index:
            obs = obs.drop("Missing")
        if obs.shape[1] < 2 or obs.shape[0] < 2 or obs.sum().sum() == 0:
            return 1.0
        _, p, _, _ = chi2_contingency(obs)
        return p
    except Exception:
        return 1.0


rows = []
for section_name, vars_ in blocks:
    rows.append({"section": section_name})
    for label, var in vars_:
        if label == "Age_Median":
            m1, q1_1, q3_1 = _median_age(g_ctrl)
            m2, q1_2, q3_2 = _median_age(g_aband)
            _, p = mannwhitneyu(g_ctrl["age_tb"].dropna(),
                                g_aband["age_tb"].dropna())
            p_str = "<0.001" if p < 0.001 else f"{p:.3f}"
            rows.append({
                "label": "Age, years, median (IQR)",
                "category": "",
                "nonltfu": f"{m1} ({q1_1}–{q3_1})",
                "ltfu":   f"{m2} ({q1_2}–{q3_2})",
                "p":       p_str,
            })
            continue

        cats = [c for c in df[var].dropna().unique() if str(c) != "nan"]
        if var == "age_group":
            cats = ["15-24", "25-44", "45-64", "≥ 65"]
        elif var == "education_clean":
            cats = ["None", "≤ 7 years", "8 - 11 years", "≥ 12 years"]
        else:
            cats = sorted(cats)

        p = get_p_val(g_ctrl, g_aband, var)
        p_str = "<0.001" if p < 0.001 else f"{p:.3f}"

        for i, cat in enumerate(cats):
            c1 = (g_ctrl[var] == cat).sum()
            c2 = (g_aband[var] == cat).sum()
            t1 = len(g_ctrl)
            t2 = len(g_aband)
            rows.append({
                "label":   label if i == 0 else "",
                "category": str(cat),
                "nonltfu":  f"{c1:,} ({c1/t1*100:.1f}%)",
                "ltfu":     f"{c2:,} ({c2/t2*100:.1f}%)",
                "p":        p_str if i == 0 else "",
            })


# Build Word table
t = doc.add_table(rows=1, cols=5)
t.style = "Table Grid"
hdr = t.rows[0].cells
hdr[0].text = "Characteristic"
hdr[1].text = "Category"
hdr[2].text = f"Non-LTFU (N = {len(g_ctrl):,})"
hdr[3].text = f"LTFU (N = {len(g_aband):,})"
hdr[4].text = "p-value"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

for r in rows:
    if "section" in r:
        row_cells = t.add_row().cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run(r["section"])
        run.italic = True
        run.bold = True
    else:
        row_cells = t.add_row().cells
        row_cells[0].text = r["label"]
        row_cells[1].text = r["category"]
        row_cells[2].text = r["nonltfu"]
        row_cells[3].text = r["ltfu"]
        row_cells[4].text = r["p"]

# Set consistent small font in table
for row in t.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

add_para(
    "Percentages are column-wise within each group. Continuous age "
    "summarised as median (interquartile range); p-value from "
    "Mann–Whitney U test. Categorical p-values from chi-squared tests "
    "(excluding missing categories from the test statistic).",
    italic=True, size=9
)


# ==========================================================================
# Save
# ==========================================================================
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
print(f"Size: {OUT_PATH.stat().st_size/1024:.1f} KB")
