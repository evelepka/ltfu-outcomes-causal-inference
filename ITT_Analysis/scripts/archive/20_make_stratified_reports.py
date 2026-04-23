import os
import pandas as pd
from docx import Document
from docx.shared import Inches

BASE_DIR = "/Users/evelynlepkadelima/Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/Abandonment Outcomes"
RESULTS_DIR = os.path.join(BASE_DIR, "Abandonment Paper/ITT_Analysis/results")

def format_pct(x):
    try:
        if pd.isna(x): return "NA"
        val = float(x)
        return f"{val:.2f}%"
    except:
        return str(x)

def build_report(report_title, out_filename, folder_name, analyses, is_long_format=False):
    doc = Document()
    doc.add_heading(report_title, 0)
    fig_dir = os.path.join(RESULTS_DIR, "manuscript figures", folder_name)

    for item in analyses:
        doc.add_heading(item["title"], level=1)
        csv_path = os.path.join(fig_dir, item["csv"])
        
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
            
            if is_long_format:
                # Pivot long format to wide format
                # columns: strata, year, mortality_pct
                df['strata'] = df['strata'].apply(lambda x: str(x).split('=')[-1] if '=' in str(x) else str(x))
                df_wide = df.pivot(index='strata', columns='year', values='mortality_pct').reset_index()
                df_wide.columns.name = None
                df_wide.rename(columns={'strata': 'Group'}, inplace=True)
                
                # Format specific years based on what's available
                years_cols = [c for c in df_wide.columns if c != 'Group']
                df_wide = df_wide[['Group'] + sorted(years_cols)]
                df = df_wide
                
            # Now df is wide format.
            for col in df.columns:
                if col != 'Group':
                    df[col] = df[col].apply(format_pct)
                    
            # Rename columns nicely
            rename_map = {}
            for col in df.columns:
                if col != 'Group':
                    if isinstance(col, (int, float)) or (isinstance(col, str) and col.replace('.','',1).isdigit()):
                        rename_map[col] = f"{float(col):.1f} Years".replace('.0', '')
                    elif col.startswith('Year_'):
                        rename_map[col] = col.replace('Year_', '') + " Years"
            df.rename(columns=rename_map, inplace=True)
            
            # Create Table
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = str(col_name)
                
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i in range(len(df.columns)):
                    row_cells[i].text = str(row.iloc[i])
                    
            doc.add_paragraph()
        else:
            doc.add_paragraph(f"Waiting for metrics table: {item['csv']}")
        
        # Add Image
        img_path = os.path.join(fig_dir, item["img"])
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))
            doc.add_paragraph()
        else:
            doc.add_paragraph(f"Waiting for figure: {item['img']}")

    out_path = os.path.join(RESULTS_DIR, out_filename)
    doc.save(out_path)
    print(f"Report saved to {out_path}")

# RETREATMENT REPORT
analyses_retr = [
    {"title": "1. Stratified by HIV Status", "csv": "retreatment_cif_hiv_values.csv", "img": "Retreatment_CIF_12y_by_HIV.png"},
    {"title": "2. Stratified by Homelessness", "csv": "retreatment_cif_homelessness_values.csv", "img": "Retreatment_CIF_12y_by_Homelessness.png"},
    {"title": "3. Stratified by Timing of Loss to Follow-up", "csv": "retreatment_cif_timing_values.csv", "img": "Retreatment_CIF_12y_by_TimingLTFU.png"}
]
build_report(
    "Stratified Retreatment Analysis (Cumulative Incidence)", 
    "Stratified_Retreatment_LTFU.docx", 
    "stratified_retreatment", 
    analyses_retr, 
    is_long_format=False
)

# MORTALITY REPORT
analyses_mort = [
    {"title": "1. Stratified by HIV Status", "csv": "ltfu_mortality_hiv_values.csv", "img": "Mortality_KM_12y_by_HIV.png"},
    {"title": "2. Stratified by Homelessness", "csv": "ltfu_mortality_homelessness_values.csv", "img": "Mortality_KM_12y_by_Homelessness.png"},
    {"title": "3. Stratified by Timing of Loss to Follow-up", "csv": "ltfu_mortality_timing_results.csv", "img": "Mortality_KM_12y_by_TimingLTFU.png"}
]
build_report(
    "Stratified Mortality Analysis (Kaplan-Meier)", 
    "Stratified_Mortality_LTFU.docx", 
    "stratified_mortality", 
    analyses_mort, 
    is_long_format=True
)

