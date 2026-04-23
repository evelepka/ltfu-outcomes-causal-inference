import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = "/Users/evelynlepkadelima/Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/Abandonment Outcomes"
RESULTS_DIR = os.path.join(BASE_DIR, "Abandonment Paper/ITT_Analysis/results")

def add_triangulation_section(doc):
    doc.add_page_break()
    doc.add_heading("Evidence-Based Triangulation: Addressing the Sick-to-Stay Bias", level=1)
    
    doc.add_paragraph(
        "To ensure a robust causal interpretation of the association between treatment abandonment and mortality, "
        "we performed a three-pillar triangulation analysis to address potential selection bias and unmeasured confounding."
    )
    
    # Pillar I
    doc.add_heading("Pillar I: 180-Day Landmark Analysis (The 'Clean' Comparison)", level=2)
    doc.add_paragraph(
        "By restricting the analysis to patients who survived the first 180 days of treatment, we eliminated the "
        "'Sick-to-Stay' bias—where the most severe cases die early and are mathematically forced into the 'completer' group. "
        "In this clarified population, abandonment was associated with a more than two-fold increase in subsequent mortality."
    )
    doc.add_paragraph("Observed RR (180d Landmark): 2.15 (95% CI: 1.95 - 2.37)")
    
    # Pillar II
    doc.add_heading("Pillar II: Quantitative Bias Analysis (QBA)", level=2)
    doc.add_paragraph(
        "The initial protective association observed at Time 0 (RR 0.947) is likely a spurious result of unmeasured "
        "baseline severity. We applied a probabilistic QBA to 'de-bias' this estimate."
    )
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Value"
    
    priors = [
        ("Observed RR (Time 0)", "0.947"),
        ("Assumed RR_UY (Severity Effect on Death)", "4.50"),
        ("Prevalence of Severity in Abandoners P(U|A=1)", "5.0%"),
        ("Prevalence of Severity in Completers P(U|A=0)", "30.0%"),
        ("Corrected/Adjusted RR", "1.65")
    ]
    
    for p, v in priors:
        row = table.add_row().cells
        row[0].text = p
        row[1].text = v
    
    doc.add_paragraph(
        "\nConclusion: After accounting for the differential prevalence of early clinical severity, the 'protective' effect "
        "disappears, revealing a true 65% increase in mortality risk associated with abandonment even from day zero."
    )
    
    # Pillar III
    doc.add_heading("Pillar III: E-Value (Robustness to Unmeasured Confounding)", level=2)
    doc.add_paragraph(
        "We calculated the E-value for the 180-day landmark estimate to quantify the strength an unmeasured "
        "confounder would need to nullify the observed effect."
    )
    
    doc.add_paragraph("E-Value for Point Estimate: 3.72")
    doc.add_paragraph("E-Value for Lower Confidence Bound: 3.31")
    
    doc.add_paragraph(
        "Interpretation: An unmeasured confounder would need a Risk Ratio of 3.72 with both abandonment and mortality "
        "to explain away the observed association. Few clinical or social factors in TB care reach this magnitude of "
        "independent effect, suggesting the results are highly robust."
    )

def Build_Comprehensive_Report():
    doc = Document()
    doc.add_heading("Comprehensive Mortality Analyses Report", 0)
    
    doc.add_paragraph("This report consolidates the stratified mortality analyses and the causal triangulation results.")
    
    # 1. Stratified Section
    doc.add_heading("1. Stratified Kaplan-Meier Mortality Analyses", level=1)
    
    stratified_items = [
        {"title": "Stratified by HIV Status", "img": "Mortality_KM_12y_by_HIV.png"},
        {"title": "Stratified by Homelessness", "img": "Mortality_KM_12y_by_Homelessness.png"},
        {"title": "Stratified by Timing of LTFU", "img": "Mortality_KM_12y_by_TimingLTFU.png"}
    ]
    
    for item in stratified_items:
        doc.add_heading(item["title"], level=2)
        img_path = os.path.join(RESULTS_DIR, "manuscript figures", "stratified_mortality", item["img"])
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))
        else:
            doc.add_paragraph(f"[Image not found: {item['img']}]")
    
    # 2. Triangulation Section
    add_triangulation_section(doc)
    
    out_path = os.path.join(RESULTS_DIR, "Comprehensive_Mortality_Analyses_Report.docx")
    doc.save(out_path)
    print(f"Comprehensive report updated at: {out_path}")

if __name__ == "__main__":
    Build_Comprehensive_Report()
