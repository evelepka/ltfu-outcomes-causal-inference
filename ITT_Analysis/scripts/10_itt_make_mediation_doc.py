import os
from docx import Document
from docx.shared import Inches

BASE_DIR = os.getcwd()
RESULTS_DIR = os.path.join(BASE_DIR, "Abandonment Paper/ITT_Analysis/results")
OUT_DOC = os.path.join(RESULTS_DIR, "Mediation_Analysis_Retreatment.docx")

doc = Document()
doc.add_heading('Mediation Analysis: The Role of Retreatment in Post-Abandonment Mortality', 0)

# Methodology Section
doc.add_heading('1. Methodology', level=1)
doc.add_paragraph('To evaluate whether subsequent TB retreatment mediates the long-term effect of initial treatment abandonment on mortality, we conducted a mediation analysis using a Time-Dependent Cox Proportional Hazards model. The cohort was restricted to individuals who survived their index treatment episode (Landmark approach at T=0).')

doc.add_paragraph('We fitted two nested survival models:')
doc.add_paragraph('1. Total Effect Model: Evaluated the overall hazard of all-cause mortality over 12 years for patients with a Loss to follow-up (LTFU) outcome compared to those with a Non-LTFU (Control) outcome. The model was adjusted for 13 baseline sociodemographic and clinical covariates (e.g., age, sex, HIV status, comorbidities, and social vulnerabilities).')
doc.add_paragraph('2. Direct Effect Model: Evaluated the same hazard but included "Retreatment" as a time-varying covariate. Patients entered the model with a retreatment status of 0, which dynamically switched to 1 on the exact date of any subsequent TB notification in the registry.')

doc.add_paragraph('The proportion of the abandonment effect mediated by retreatment was calculated based on the reduction in the log-Hazard Ratio between the two models. This approach isolates the direct mortality risk of abandonment from the indirect risk strictly mediated by falling into a retreatment cycle.')

# Results Section
doc.add_heading('2. Results', level=1)
doc.add_paragraph('The analysis demonstrated that retreatment is a substantial mediator of post-abandonment mortality.')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Effect Pathway'
hdr[1].text = 'Hazard Ratio (HR)'
hdr[2].text = 'Interpretation'

row1 = table.add_row().cells
row1[0].text = 'Total Effect (Abandonment -> Death)'
row1[1].text = '1.74'
row1[2].text = 'Patients who abandon treatment have a 74% higher overall mortality risk in the subsequent decade.'

row2 = table.add_row().cells
row2[0].text = 'Direct Effect (Abandonment -> Death)'
row2[1].text = '1.25'
row2[2].text = 'Removing the influence of retreatment (holding it constant), the excess mortality risk drops to 25%.'

row3 = table.add_row().cells
row3[0].text = 'Mediator Effect (Retreatment -> Death)'
row3[1].text = '2.54'
row3[2].text = 'The act of returning to the system for retreatment independently carries a 2.54x mortality penalty.'

doc.add_paragraph('\nConclusion:', style='Intense Quote')
p = doc.add_paragraph('Based on the models, approximately ')
p.add_run('59.7%').bold = True
p.add_run(' of the total mortality effect of abandonment is mediated by subsequent retreatment. This indicates that while initial abandonment leaves some residual long-term damage, the majority of its lethal consequence occurs because it forces the patient to face a recurrent, highly lethal TB retreatment episode.')

doc.save(OUT_DOC)
print(f"Mediation Word report saved to {OUT_DOC}")
