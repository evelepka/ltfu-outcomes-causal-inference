import pandas as pd
from docx import Document
from docx.shared import Inches
import os

# Paths
BASE_DIR = os.getcwd()
RESULTS_DIR = os.path.join(BASE_DIR, "Abandonment Paper/ITT_Analysis/results")
CSV_PATH = os.path.join(RESULTS_DIR, "g_formula_mortality_12y_results.csv")
PNG_PATH = os.path.join(RESULTS_DIR, "g_formula_mortality_12y.png")
OUT_DOC = os.path.join(RESULTS_DIR, "G_Formula_Results_12y.docx")

# 1. Load Data
df = pd.read_csv(CSV_PATH)

# 2. Create Document
doc = Document()
doc.add_heading('G-Formula Causal Analysis (12-Year Mortality)', 0)

doc.add_paragraph('This analysis estimates the counterfactual cumulative mortality risk over a 12-year (144 months) follow-up period, comparing "Loss to follow-up (LTFU)" vs. "Non-LTFU (Control)".')

# 3. Add Table
doc.add_heading('Causal Risk Estimates', level=1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = 'Time'
hdr[1].text = 'Non-LTFU (Control)'
hdr[2].text = 'Loss to follow-up'
hdr[3].text = 'Risk Ratio'
hdr[4].text = 'Excess Mortality (%)'

for _, row in df.iterrows():
    r = table.add_row().cells
    r[0].text = str(row['Time'])
    r[1].text = f"{row['CI_Non_LTFU']*100:.2f}%"
    r[2].text = f"{row['CI_LTFU']*100:.2f}%"
    r[3].text = f"{row['Risk_Ratio']:.2f}"
    r[4].text = f"{row['Risk_Diff']*100:+.2f}%"

# 4. Add Image
doc.add_heading('Survival Curves', level=1)
doc.add_picture(PNG_PATH, width=Inches(6))

# 5. Save
doc.save(OUT_DOC)
print(f"G-Formula Word report saved to {OUT_DOC}")
