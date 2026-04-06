import pandas as pd
from docx import Document
import os
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
IN_CSV = os.path.join(BASE_DIR, "ITT_Analysis/results/multivariable_results_mi_cc.csv")
OUT_DOC = os.path.join(BASE_DIR, "ITT_Analysis/results/Table2_ITT.docx")

df = pd.read_csv(IN_CSV)

def get_val(model_n, term_name):
    row = df[(df["model"] == model_n) & (df["term"] == term_name)]
    if row.empty: return ""
    return row.iloc[0]["estimate"]

structure = [
    ("Socio-demographic", [
        ("Age Group (ref: 15–24 years)", ""),
        ("25–44 years", "age_group25-44"),
        ("45–64 years", "age_group45-64"),
        ("≥65 years", "age_group≥65"),
        ("Male Sex", "sexMale"),
        ("Race (ref: White)", ""),
        ("Black or Mixed", "race_cleanBlack or Mixed"),
        ("Other", "race_cleanOther"),
        ("Education Level (ref: ≥12 years)", ""),
        ("8–11 years", "edu_clean8 - 11 years"),
        ("≤7 years", "edu_clean≤ 7 years"),
        ("No education", "edu_cleanNone")
    ]),
    ("Clinical Comorbidities", [
        ("HIV/AIDS status (ref: Negative)", ""),
        ("Positive", "hiv_aidsPositive"),
        ("Diabetes", "diabetesYes")
    ]),
    ("Behavioral", [
        ("Alcohol Use", "alcoholYes"),
        ("Drug Use", "drug_useYes")
    ]),
    ("Social Vulnerabilities", [
        ("Incarcerated", "incarceratedYes"),
        ("Homelessness", "homelessnessYes")
    ]),
    ("Disease and Treatment-Related", [
        ("Hospital Admission", "hosp_admissionYes"),
        ("Clinical form (ref: Pulmonary)", ""),
        ("Extrapulmonary", "clinical_cleanExtrapulmonary"),
        ("Mixed/Disseminated", "clinical_cleanPulmonary and Extrapulmonary or disseminated"),
        ("Direct Observed Therapy (DOT)", "dot_statusYes"),
        ("Loss to follow-up duration (ref: ≥ 4 months)", ""),
        ("< 2 months", "tx_month_grp< 2 months"),
        ("2 to <4 months", "tx_month_grp2 to <4 months")
    ])
]

def create_mv_table(doc, title, suffix_adj):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'

    hdr1 = table.rows[0].cells
    hdr1[0].text = "Characteristics"
    hdr1[1].text = "Retreatment (SHR)"
    hdr1[1].merge(hdr1[2]) 
    hdr1[3].text = "Mortality (HR)"
    hdr1[3].merge(hdr1[4]) 

    hdr2 = table.rows[1].cells
    hdr2[1].text = "Unadjusted"
    hdr2[2].text = "Adjusted"
    hdr2[3].text = "Unadjusted"
    hdr2[4].text = "Adjusted"

    header_keywords = ["(ref:", "Male Sex", "Positive", "Diabetes", "Alcohol Use", "Drug Use", "Incarcerated", "Homelessness", "Hospital Admission", "Direct Observed Therapy (DOT)", "< 2 months"]

    for category_name, vars in structure:
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run(category_name)
        run.italic = True
        
        for label, term in vars:
            row = table.add_row().cells
            prefix = "    " if term != "" and not any(k in label for k in header_keywords) else ""
            row[0].text = prefix + label
            if term:
                row[1].text = get_val("FG_Retr_Unadjusted", term)
                row[2].text = get_val(f"FG_Retr_Adjusted_{suffix_adj}", term)
                row[3].text = get_val("Cox_Death_Unadjusted", term)
                row[4].text = get_val(f"Cox_Death_Adjusted_{suffix_adj}", term)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = 0
    
    doc.add_paragraph("\n") # Spacer

# Get actual counts from result file or use the ones from the run
doc = Document()
doc.add_heading("Table 2: Multivariable Results (ITT Cohort)", 0)

create_mv_table(doc, "Table 2a: Multiple Imputation (N=21,619 Loss to follow-up)", "MI")
create_mv_table(doc, "Table 2b: Complete Case Analysis (N=11,864 Loss to follow-up)", "CC")

doc.save(OUT_DOC)
print(f"Table 2 saved to {OUT_DOC}")

