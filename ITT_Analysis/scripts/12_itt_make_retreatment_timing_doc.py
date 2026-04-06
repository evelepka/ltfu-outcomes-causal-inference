import os
from docx import Document
from docx.shared import Inches

BASE_DIR = os.getcwd()
RESULTS_DIR = os.path.join(BASE_DIR, "Abandonment Paper/ITT_Analysis/results")
OUT_DOC = os.path.join(RESULTS_DIR, "Retreatment_Timing_Analysis.docx")

doc = Document()
doc.add_heading('Analysis of Retreatment Timing and Post-Retreatment Mortality', 0)

# Methodology Section
doc.add_heading('1. Methodology', level=1)
doc.add_paragraph('To answer the clinical question of whether the specific "month of retreatment" acts as an independent predictor of post-retreatment mortality, we conducted a specialized survival analysis.')
doc.add_paragraph('Cohort Restriction: The analysis was strictly limited to patients who re-entered the TB registry for a subsequent treatment episode (evaluating only "retreaters").')
doc.add_paragraph('Baseline (T=0): The survival clock was initiated on the exact date the patient started their retreatment (not their index treatment).')
doc.add_paragraph('Exposure Variable: The delay between the completion/abandonment of the first episode and the initiation of the retreatment episode. This "Time to Retreatment" was categorized into four clinical intervals: Early (0-6 months), Intermediate (6-12 months), Late (1-3 years), and Very Late (>3 years).')
doc.add_paragraph('Statistical Model: A Multivariable Cox Proportional Hazards model was fitted to estimate the risk of all-cause mortality following retreatment. The model was adjusted for the 13 baseline sociodemographic and clinical covariates, as well as the clinical outcome of the patient\'s first TB episode (Abandonment vs. Cure).')

# Results Section
doc.add_heading('2. Results', level=1)
doc.add_paragraph('The Multivariable Cox model revealed that the chronological delay to retreatment does not significantly impact survival, but the origin of the retreatment does.')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Predictor'
hdr[1].text = 'Hazard Ratio (HR)'
hdr[2].text = 'P-value'
hdr[3].text = 'Interpretation'

row1 = table.add_row().cells
row1[0].text = 'Timing: Very Late (>3 years)'
row1[1].text = 'Reference'
row1[2].text = '-'
row1[3].text = 'Baseline comparison group.'

row2 = table.add_row().cells
row2[0].text = 'Timing: Late (1-3 years)'
row2[1].text = '1.29'
row2[2].text = '0.109'
row2[3].text = 'No statistically significant difference compared to Very Late returns.'

row3 = table.add_row().cells
row3[0].text = 'Timing: Intermediate (6-12m)'
row3[1].text = '1.16'
row3[2].text = '0.366'
row3[3].text = 'No statistically significant difference compared to Very Late returns.'

row4 = table.add_row().cells
row4[0].text = 'Timing: Early (0-6m)'
row4[1].text = '1.16'
row4[2].text = '0.299'
row4[3].text = 'No statistically significant difference compared to Very Late returns.'

row5 = table.add_row().cells
row5[0].text = 'Prior Outcome: Abandonment'
row5[1].text = '1.22'
row5[2].text = '0.034'
row5[3].text = 'Patients returning due to prior Abandonment have a 22.3% higher mortality risk during/after retreatment than those returning due to true Relapse.'

doc.add_paragraph('\nClinical Conclusion:', style='Intense Quote')
p = doc.add_paragraph('The chronological delay to retreatment does not independently drive prognostic mortality risk. ')
p.add_run('Early returns are not inherently more fatal than very late returns.').bold = True
p.add_run(' Instead, the critical determinant of survival is ')
p.add_run('why').italic = True
p.add_run(' the patient is returning. Patients forced to re-enter the system because they abandoned their first treatment carry a significant 22% mortality penalty compared to patients who return due to a biological relapse after an initial cure, establishing abandonment as a persistent marker of vulnerability across subsequent episodes.')

doc.save(OUT_DOC)
print(f"Retreatment Timing Word report saved to {OUT_DOC}")
