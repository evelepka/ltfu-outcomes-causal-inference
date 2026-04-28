"""
Build a short Word document with the calendar-year and drug-resistance
sensitivity-analysis paragraphs, for later inclusion in the manuscript.
Saved to Drafts/ in the project Google Drive.
"""

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _find_project_root() -> Path:
    if os.environ.get("TB_ABANDONMENT_ROOT"):
        p = Path(os.environ["TB_ABANDONMENT_ROOT"]).expanduser()
        if p.exists():
            return p
    for c in [
        Path.home() / "Library/CloudStorage/GoogleDrive-jasonandr@gmail.com/My Drive/Abandonment Paper",
        Path.home() / "Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/Abandonment Outcomes/Abandonment Paper",
    ]:
        if c.exists():
            return c
    return Path(__file__).resolve().parents[2]


BASE = _find_project_root()
OUT_DIR = BASE / "Drafts"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / f"Sensitivity_text_year_dr_{date.today().isoformat()}.docx"

doc = Document()

# Set a reasonable default font size
style = doc.styles["Normal"]
style.font.size = Pt(11)

doc.add_heading("Sensitivity-analysis text — calendar-year trends and drug-resistance status", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "Short paragraphs intended for the main-text Methods/Results sensitivity "
    "section. Companion supplement tables: S_year_ltfu, S_year_outcomes_ltfu, "
    "and S_dr_status (Tables/ folder). Numbers reflect the ITT cohort (N=172,463; "
    "21,619 LTFU, 150,844 Non-LTFU) and the per-patient drug-resistance "
    "derivation described in Table S_dr_status."
)
run.italic = True
run.font.size = Pt(10)

doc.add_heading("Calendar-year sensitivity analysis", level=2)
doc.add_paragraph(
    "In a calendar-year sensitivity analysis (Table S_year_outcomes), 1- and "
    "2-year mortality (approximately 2–3% and 4–5%, respectively) and "
    "retreatment cumulative incidence (approximately 30% and 37%) among LTFU "
    "patients were stable across 2013–2022, with no detectable change during "
    "the COVID-19 period, despite a concurrent rise in LTFU incidence from "
    "approximately 10% in 2013–2018 to approximately 15% in 2020–2022 "
    "(Table S_year_ltfu). The 2023 cohort was excluded from the calendar-year "
    "outcome analysis because it is right-truncated by the inclusion criterion "
    "(end_date ≤ 2023-12-31), which disproportionately removes patients with "
    "longer treatment durations."
)

doc.add_heading("Drug-resistance sensitivity analysis", level=2)
doc.add_paragraph(
    "In a sensitivity analysis using a multi-source drug-resistance "
    "classification that combined Xpert MTB/RIF results, the SINAN drug-"
    "susceptibility-testing summary, and drug-specific rifampin and isoniazid "
    "DST results (Table S_dr_status), LTFU rates were highest among patients "
    "with rifampin-resistant or multidrug-resistant TB (20.4%, 95% CI 18.5–22.4) "
    "compared with patients with isoniazid-monoresistant disease (15.6%, "
    "13.4–18.1) and drug-sensitive disease (16.4%, 16.1–16.7). LTFU rates "
    "were lowest in the Not Evaluated group (9.7%), which is consistent with "
    "differential ascertainment of DST (more frequently performed in patients "
    "with smear-positive pulmonary disease and longer index episodes) rather "
    "than a true protective effect of having no DST result."
)

doc.save(OUT)
print(f"Wrote {OUT}")
