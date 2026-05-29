import pandas as pd
import numpy as np
from scipy.stats import chi2_contingency, mannwhitneyu
import os
from docx import Document
from docx.shared import Pt
import warnings
warnings.filterwarnings("ignore")

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DATA_CSV = os.path.join(BASE_DIR, "ITT_Analysis/data/itt_cohort.csv")
OUT_DOC = os.path.join(BASE_DIR, "ITT_Analysis/results/Table1_ITT.docx")

df = pd.read_csv(DATA_CSV, low_memory=False)

# The variables are already pre-cleaned in itt_cohort.csv (which mirrors analysis_ready_cohort.csv).
# However, we still need to remap a few categories for Table 1 presentation directly.

df["sex"] = df["sex"].fillna("Missing")
df["age_group"] = pd.cut(df["age_tb"], bins=[14, 24, 44, 64, 150], labels=["15-24", "25-44", "45-64", "≥ 65"])

df["hiv_aids"] = df["hiv_aids"].fillna("Missing")
# 2026-05-29: cohort uses these column names; map old→new
df["incarceration"] = df["incarcerated"].fillna("Missing")
df["homelessness"] = df["homelessness"].fillna("Missing")
df["supervised_therapy"] = df["dot_status"].fillna("Missing")
df["alcohol"] = df["alcohol"].fillna("Missing")
df["drug_use_stat"] = df["drug_use"].fillna("Missing")
df["tobacco_stat"] = df["tobacco_use"].fillna("Missing")
df["diabetes"] = df["diabetes"].fillna("Missing")
df["mental"] = df["mental_health"].fillna("Missing")
df["immuno"] = df["other_immuno_condition"].fillna("Missing")

# clinical_clean is already cleaned in the new cohort CSV
df["clinical_clean"] = df["clinical_clean"].fillna("Missing")

df["diagnosis_clean"] = df["diagnosis_setting"].fillna("Missing")
df["lab_confirmed_stat"] = df["lab_confirmed_stat"].fillna("Missing")
df["hosp_admission_stat"] = df["hosp_admission"].fillna("Missing")

# race_clean and edu_clean are already cleaned in the new cohort CSV
df["race_clean"] = df["race_clean"].fillna("Missing")
df["education_clean"] = df["edu_clean"].fillna("Missing")

# Split Groups
g_aband = df[df["itt_group"] == "Loss to follow-up"]
g_ctrl = df[df["itt_group"] == "Non-LTFU"]

blocks = [
    ("Socio-demographic", [
        ("Age_Median", "age_tb"), # Special case handled in loop
        ("Age Group", "age_group"),
        ("Sex", "sex"),
        ("Race", "race_clean"),
        ("Education Level", "education_clean")
    ]),
    ("Clinical Comorbidities", [
        ("HIV/AIDS", "hiv_aids"),
        ("Diabetes", "diabetes"),
        ("Other immunosuppressive conditions than HIV", "immuno"),
        ("Mental Health Issues", "mental")
    ]),
    ("Behavioral", [
        ("Alcohol Use", "alcohol"),
        ("Drug Use", "drug_use_stat"),
        ("Tobacco Use", "tobacco_stat")
    ]),
    ("Social Vulnerabilities", [
        ("Incarcerated", "incarceration"),
        ("Homelessness", "homelessness")
    ]),
    ("Disease and Treatment-Related", [
        ("Diagnosis Setting", "diagnosis_clean"),
        ("Hospital Admission", "hosp_admission_stat"),
        ("Laboratory Confirmed Status", "lab_confirmed_stat"),
        ("Clinical Classification", "clinical_clean"),
        ("Direct Observed Therapy (DOT)", "supervised_therapy")
    ])
]

results = []

def get_p_val(df1, df2, var):
    try:
        obs = pd.concat([df1[var].value_counts(), df2[var].value_counts()], axis=1).fillna(0)
        if 'Missing' in obs.index: obs = obs.drop('Missing')
        if obs.shape[1] < 2 or obs.shape[0] < 2 or obs.sum().sum() == 0: return 1.0
        _, p, _, _ = chi2_contingency(obs)
        return p
    except:
        return 1.0

for section_name, vars_in_section in blocks:
    results.append({
        "Variable": section_name, "Category": "", "Total": "", "Non-LTFU": "", "Loss to follow-up": "", "P-value": ""
    })
    
    for label, var in vars_in_section:
        if label == "Age_Median":
            mt, q1_t, q3_t = df['age_tb'].median(), df['age_tb'].quantile(0.25), df['age_tb'].quantile(0.75)
            m1, q1_1, q3_1 = g_ctrl['age_tb'].median(), g_ctrl['age_tb'].quantile(0.25), g_ctrl['age_tb'].quantile(0.75)
            m2, q1_2, q3_2 = g_aband['age_tb'].median(), g_aband['age_tb'].quantile(0.25), g_aband['age_tb'].quantile(0.75)
            _, p_age = mannwhitneyu(g_ctrl['age_tb'].dropna(), g_aband['age_tb'].dropna())
            results.append({"Variable": "Age (years, median/IQR)", "Category": "",
                            "Total": f"{int(mt)} ({int(q1_t)}–{int(q3_t)})",
                            "Non-LTFU": f"{int(m1)} ({int(q1_1)}–{int(q3_1)})",
                            "Loss to follow-up": f"{int(m2)} ({int(q1_2)}–{int(q3_2)})",
                            "P-value": f"{p_age:.3f}" if p_age >= 0.001 else "<0.001"})
            continue
            
        cats = [c for c in df[var].dropna().unique() if str(c) != "nan"]
        if var == "age_group":
            cats = ["15-24", "25-44", "45-64", "≥ 65"]
        elif var == "education_clean":
            cats = ["None", "≤ 7 years", "8 - 11 years", "≥ 12 years"]
        else:
            cats = sorted(cats)
            
        p_val = get_p_val(g_ctrl, g_aband, var)
        p_str = f"{p_val:.3f}" if p_val >= 0.001 else "<0.001"
        
        for i, cat in enumerate(cats):
            ct, tt = (df[var] == cat).sum(), len(df)
            c1, t1 = (g_ctrl[var] == cat).sum(), len(g_ctrl)
            c2, t2 = (g_aband[var] == cat).sum(), len(g_aband)
            
            row = {
                "Variable": label if i == 0 else "",
                "Category": str(cat),
                "Total": f"{ct:,} ({ct/tt*100:.1f}%)" if tt > 0 else "0 (0%)",
                "Non-LTFU": f"{c1:,} ({c1/t1*100:.1f}%)" if t1 > 0 else "0 (0%)",
                "Loss to follow-up": f"{c2:,} ({c2/t2*100:.1f}%)" if t2 > 0 else "0 (0%)",
                "P-value": p_str if i == 0 else ""
            }
            results.append(row)

# Create docx
doc = Document()
doc.add_heading(f"Table 1: Baseline Characteristics (N={len(df):,})", level=1)
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Characteristics"
hdr_cells[1].text = "Category"
hdr_cells[2].text = f"Total (N={len(df):,})"
hdr_cells[3].text = f"Non-LTFU (N={len(g_ctrl):,})"
hdr_cells[4].text = f"Loss to follow-up (N={len(g_aband):,})"
hdr_cells[5].text = "p-value"
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

for r in results:
    row_cells = table.add_row().cells
    if r["Category"] == "" and r["Total"] == "":
        p = row_cells[0].paragraphs[0]
        run = p.add_run(r["Variable"])
        run.italic = True
    else:
        row_cells[0].text = r["Variable"]
        row_cells[1].text = r["Category"]
        row_cells[2].text = r["Total"]
        row_cells[3].text = r["Non-LTFU"]
        row_cells[4].text = r["Loss to follow-up"]
        row_cells[5].text = r["P-value"]

doc.save(OUT_DOC)
print(f"Table 1 saved to {OUT_DOC}")
