import pandas as pd
import numpy as np
import os
from pathlib import Path
from docx import Document

# ----------------------------------------------------------------------------
# Project root resolution
# ----------------------------------------------------------------------------
# The raw SINAN data and cohort outputs live in Google Drive, not in git.
# Resolve in this order:
#   1. TB_ABANDONMENT_ROOT environment variable (set this for non-standard setups)
#   2. Known Google Drive mounts for current collaborators
#   3. Script-relative fallback (script_dir/../.. -> works if Data/ is co-located)
def _find_project_root():
    if os.environ.get("TB_ABANDONMENT_ROOT"):
        p = Path(os.environ["TB_ABANDONMENT_ROOT"]).expanduser()
        if p.exists():
            return p
    candidates = [
        Path.home() / "Library/CloudStorage/GoogleDrive-jasonandr@gmail.com/My Drive/Abandonment Paper",
        Path.home() / "Library/CloudStorage/GoogleDrive-evelynlepka@gmail.com/My Drive/Abandonment Outcomes/Abandonment Paper",
    ]
    for c in candidates:
        if c.exists():
            return c
    # Last resort: repo-relative (only works if Data/ is co-located with ITT_Analysis/)
    return Path(__file__).resolve().parents[2]

BASE_DIR = _find_project_root()
FINAL_DATA = BASE_DIR / "Data" / "Final_table_cleaned.csv"
OUT_CSV = BASE_DIR / "ITT_Analysis" / "data" / "itt_cohort.csv"
OUT_DOC = BASE_DIR / "ITT_Analysis" / "results" / "Inclusion_Exclusion_ITT.docx"
OUT_FLOWCHART = BASE_DIR / "Data" / "exclusion_flowchart.csv"

OUT_CSV.parent.mkdir(parents=True, exist_ok=True)
OUT_DOC.parent.mkdir(parents=True, exist_ok=True)

print(f"Base Dir: {BASE_DIR}")
if not FINAL_DATA.exists():
    raise FileNotFoundError(
        f"Final_table_cleaned.csv not found at {FINAL_DATA}. "
        f"Set TB_ABANDONMENT_ROOT or mount Google Drive."
    )

# 1. Load the full dataset
print("Loading cleaned dataset...")
df = pd.read_csv(FINAL_DATA, low_memory=False)
for col in ["notification_date", "end_date", "dod", "tx_start"]:
    df[col] = pd.to_datetime(df[col], errors="coerce")

# 2. Ascertain ALL deaths per individual (Source A: dod, Source B: case_outcome='Obito')
print("Aggregating overall deaths across all records...")
# Source A: dod field
dod_src_a = df.dropna(subset=["dod"]).groupby("sinan_clean")["dod"].max().reset_index().rename(columns={"dod": "death_src_a"})
# Source B: Obito outcome
OBITO_OUTCOMES = {"Obito TB", "Obito NTB"}
obito_rows = df[df["case_outcome"].isin(OBITO_OUTCOMES)].dropna(subset=["end_date"])
dod_src_b = obito_rows.groupby("sinan_clean")["end_date"].max().reset_index().rename(columns={"end_date": "death_src_b"})
dod_all = dod_src_a.merge(dod_src_b, on="sinan_clean", how="outer")
dod_all["death_date_comprehensive"] = dod_all[["death_src_a", "death_src_b"]].max(axis=1)

# 3. Separate Treatment and Control Candidates (First Novo instance)
ABANDON_OUTCOMES = {"Abandono", "Abandono Primario", "Faltoso"}

# Load diagnostic_date for all IDs to use as proxy
print("Loading diagnostic dates for proxying...")
diag_df = df[["sinan_clean", "diagnostic_date"]].dropna().groupby("sinan_clean")["diagnostic_date"].min().reset_index()

# For the LTFU group: First "Novo" episode ending in abandonment
print("Selecting First Abandonment episodes (Novo)...")
# Excluding "Mud Diag", empty/null, and blank outcomes — none of these represent a valid TB episode conclusion
all_novo = df[df["case_type"].str.strip().str.title() == "Novo"]
all_novo_valid = all_novo[
    all_novo["case_outcome"].notna() &
    (all_novo["case_outcome"].str.strip() != "") &
    (all_novo["case_outcome"] != "Mud Diag")
]

# To avoid overlap and ensure we pick the *absolute first* episode correctly, 
# we take the first Novo per person among ANY outcome (except Mud Diag), then split.
first_novo = all_novo_valid.sort_values("end_date").groupby("sinan_clean", as_index=False).first().copy()

# Outcomes to exclude entirely from the ITT cohort (not abandonment, not valid control)
TRANSFER_OUTCOMES = {"Transf Outro Municipio", "Transf Outro Estado/Pais"}

# Classify based on index outcome
# Non-LTFU: only valid definitive outcomes (excludes transfers)
first_novo["itt_group"] = np.where(
    first_novo["case_outcome"].isin(ABANDON_OUTCOMES), "Loss to follow-up",
    np.where(first_novo["case_outcome"].isin(TRANSFER_OUTCOMES), "Excluded_Transfer", "Non-LTFU")
)

# Remove transfers from the cohort entirely
first_novo = first_novo[first_novo["itt_group"] != "Excluded_Transfer"]

# Use diagnostic_date already present in first_novo
itt_cohort = first_novo.copy()
for col in ["diagnostic_date"]:
    itt_cohort[col] = pd.to_datetime(itt_cohort[col], errors="coerce")

# 4. Apply Study Metadata and Common Filters
itt_cohort["age_group"] = pd.cut(itt_cohort["age_tb"], bins=[15, 25, 45, 65, 120], labels=["15-24", "25-44", "45-64", "65+"], right=False)

print("Applying study filters...")
attrition = []
attrition.append(("Initial Candidates (First Novo Episode per Individual)", len(itt_cohort)))

# Filter for Age >= 15
itt_cohort = itt_cohort[itt_cohort["age_tb"] >= 15]
attrition.append(("Exclude: Age < 15 years", len(itt_cohort)))

# Filter for Window (end_date 2013-2023)
itt_cohort = itt_cohort.dropna(subset=["end_date"])
itt_cohort = itt_cohort[(itt_cohort["end_date"] >= pd.Timestamp("2013-01-01")) & (itt_cohort["end_date"] <= pd.Timestamp("2023-12-31"))]
attrition.append(("Exclude: Treatment end date outside 2013-2023", len(itt_cohort)))

# --- PROXY START DATE LOGIC ---
# This is crucial for avoiding data loss when tx_start is missing.
# Proxy sequence: treatment start -> diagnostic date -> notification date
itt_cohort["best_start"] = itt_cohort["tx_start"].fillna(itt_cohort["diagnostic_date"]).fillna(itt_cohort["notification_date"])
itt_cohort = itt_cohort.merge(dod_all[["sinan_clean", "death_date_comprehensive"]], on="sinan_clean", how="left")

# Exclude candidates who died ON or BEFORE treatment/proxy start (Academic Rigor)
# This handles the "Immortal Time" at the beginning of the treatment.
pre_tx_death = (itt_cohort["death_date_comprehensive"].notna()) & (itt_cohort["death_date_comprehensive"] <= itt_cohort["best_start"])
itt_cohort = itt_cohort[~pre_tx_death]
attrition.append(("Exclude: Death on or before treatment start/proxy date", len(itt_cohort)))

# Filter valid end dates (if tx_start is present, must be <= end_date)
# (Used but also handled by tx_duration logic downstream)
invalid_mask = (itt_cohort["tx_start"].notna()) & (itt_cohort["end_date"] < itt_cohort["tx_start"])
itt_cohort = itt_cohort[~invalid_mask]
attrition.append(("Exclude: Invalid dates (end_date < tx_start)", len(itt_cohort)))

print(f"Final ITT Cohort Size: {len(itt_cohort):,}")
print(itt_cohort["itt_group"].value_counts())

# 5. Ascertain Outcomes
print("Ascertaining subsequent outcomes...")
CENSOR_DATE = pd.Timestamp("2024-12-31")

# Re-notification: Merge with original df to find any notification_date > index end_date
sub_events = df.merge(itt_cohort[["sinan_clean", "end_date"]], on="sinan_clean", suffixes=("", "_index"))
sub_events = sub_events[sub_events["notification_date"] > sub_events["end_date_index"]]
next_notif = sub_events.groupby("sinan_clean")["notification_date"].min().reset_index().rename(columns={"notification_date": "next_notif_date"})
itt_cohort = itt_cohort.merge(next_notif, on="sinan_clean", how="left")
itt_cohort.loc[itt_cohort["next_notif_date"] > CENSOR_DATE, "next_notif_date"] = pd.NaT

# Death outcome: Already merged death_date_comprehensive
# Using >= to include deaths at the exact same day of abandonment
itt_cohort["death_date"] = pd.to_datetime(np.where(itt_cohort["death_date_comprehensive"] >= itt_cohort["end_date"], 
                                                  itt_cohort["death_date_comprehensive"], pd.NaT))
itt_cohort.loc[itt_cohort["death_date"] > CENSOR_DATE, "death_date"] = pd.NaT

# assign_renotif logic (Competing Risks)
def assign_renotif_itt(row):
    t_n = row["next_notif_date"]
    t_d = row["death_date"]
    has_n, has_d = pd.notna(t_n), pd.notna(t_d)
    if has_n and has_d:
        if t_n <= t_d: return (1, (t_n - row["end_date"]).days / 365.25)
        else:          return (2, (t_d - row["end_date"]).days / 365.25)
    elif has_n: return (1, (t_n - row["end_date"]).days / 365.25)
    elif has_d: return (2, (t_d - row["end_date"]).days / 365.25)
    else:       return (0, (CENSOR_DATE - row["end_date"]).days / 365.25)

rn_out = itt_cohort.apply(assign_renotif_itt, axis=1, result_type="expand")
itt_cohort["event_rn"], itt_cohort["time_rn"] = rn_out[0], rn_out[1]

# assign_death logic
def assign_death_itt(row):
    t_d = row["death_date"]
    if pd.notna(t_d): return (1, (t_d - row["end_date"]).days / 365.25)
    else:           return (0, (CENSOR_DATE - row["end_date"]).days / 365.25)

d_out = itt_cohort.apply(assign_death_itt, axis=1, result_type="expand")
itt_cohort["event_d"], itt_cohort["time_d"] = d_out[0], d_out[1]

# New: Survival times from tx_start (for G-formula consistency)
def assign_death_tx(row):
    t_d = row["death_date"]
    if pd.notna(t_d): return (t_d - row["best_start"]).days / 365.25
    else:           return (CENSOR_DATE - row["best_start"]).days / 365.25
itt_cohort["time_d_tx"] = itt_cohort.apply(assign_death_tx, axis=1)

def assign_renotif_tx(row):
    t_n = row["next_notif_date"]
    t_d = row["death_date"]
    has_n, has_d = pd.notna(t_n), pd.notna(t_d)
    if has_n and has_d:
        return ((min(t_n, t_d) - row["best_start"]).days / 365.25)
    elif has_n: return (t_n - row["best_start"]).days / 365.25
    elif has_d: return (t_d - row["best_start"]).days / 365.25
    else:       return (CENSOR_DATE - row["best_start"]).days / 365.25
itt_cohort["time_rn_tx"] = itt_cohort.apply(assign_renotif_tx, axis=1)

# 6. Variable Cleaning (Harmonize covariates)
def yn(x):
    s = str(x).strip().upper()
    if s == "S": return "Yes"
    if s == "N": return "No"
    return np.nan

itt_cohort["hiv_aids"] = np.where((itt_cohort["hiv"] == "Pos") | (itt_cohort["aids"] == "S"), "Positive", "Negative")
itt_cohort.loc[itt_cohort["hiv"].isna() & itt_cohort["aids"].isna(), "hiv_aids"] = np.nan

itt_cohort["dot_status"] = itt_cohort["tx_administration_type"].apply(lambda x: "Yes" if "Supervisionado" in str(x) else ("No" if pd.notna(x) else np.nan))
itt_cohort["incarcerated"] = itt_cohort["address_type"].apply(lambda x: "Yes" if "DETENTO" in str(x).upper() else ("No" if pd.notna(x) else np.nan))
itt_cohort["homelessness"] = itt_cohort["address_type"].apply(lambda x: "Yes" if "SEM RESIDENCIA FIXA" in str(x).upper() else ("No" if pd.notna(x) else np.nan))

itt_cohort["alcohol"] = itt_cohort["alcoholism"].apply(yn)
itt_cohort["drug_use"] = itt_cohort["drug_use"].apply(yn)
itt_cohort["mental_health"] = itt_cohort["mental_issue"].apply(yn)
itt_cohort["diabetes"] = itt_cohort["diabetes"].apply(yn)
itt_cohort["tobacco_use"] = itt_cohort["tobacco_use"].apply(yn)
itt_cohort["other_immuno_condition"] = itt_cohort["other_immuno_condition"].apply(yn)
itt_cohort["hosp_admission"] = itt_cohort["hosp_admission"].apply(yn)

itt_cohort["sex"] = itt_cohort["sex"].replace({"M": "Male", "F": "Female"})
itt_cohort["race_clean"] = itt_cohort["race"].replace({"Amarelo": "Other", "Indigena": "Other", "Pardo": "Black or Mixed", "Preto": "Black or Mixed", "Branco": "White", "Ignorado": np.nan})
itt_cohort["edu_clean"] = itt_cohort["education"].replace({"Nenhuma": "None", "De 1 a 3 anos": "≤ 7 years", "De 4 a 7 anos": "≤ 7 years", "De 8 a 11 anos": "8 - 11 years", "De 12 a 14 anos": "≥ 12 years", "15 anos e mais": "≥ 12 years", "Ignorado": np.nan})

itt_cohort["clinical_clean"] = itt_cohort["clinical_classif"].replace({
    "Pul": "Pulmonary", "Ext": "Extrapulmonary", "P+E": "Pulmonary and Extrapulmonary or disseminated", "Dissem": "Pulmonary and Extrapulmonary or disseminated"
})
itt_cohort["diagnosis_setting"] = itt_cohort["disease_discovery"].replace({
    "Demanda Ambulatorial": "Outpatient", "Urgencia / Emergencia": "Emergency / Inpatient", "Elucidacao Diagn. em Internacao": "Emergency / Inpatient",
    "Busca Ativa em Instituicao": "Active finding in institution", "Busca Ativa na Comunidade": "Active finding in community", "Investigacao de Contatos": "Contact investigation", "S/inf": np.nan
})
itt_cohort["lab_confirmed_stat"] = itt_cohort["lab_confirmed"].apply(lambda x: "Yes" if x == 1 else ("No" if x == 0 else np.nan))

# New Bacteriological Covariates
itt_cohort["bac1_clean"] = itt_cohort["bac1"].replace({
    "Pos": "Positive", "Neg": "Negative", "N/realiz": "Not Evaluated", "S/inf": "Not Evaluated", "And": "Not Evaluated"
}).fillna("Not Evaluated")

itt_cohort["sputum_culture_clean"] = itt_cohort["sputum_culture"].replace({
    "Pos": "Positive", "Neg": "Negative", "N/realiz": "Not Evaluated", "S/inf": "Not Evaluated", "And": "Not Evaluated"
}).fillna("Not Evaluated")

# For resistance, any recorded state means drug susceptibility testing was performed, NA means not evaluated
def parse_res(r):
    if pd.isna(r): return "Not Evaluated"
    r_str = str(r).upper().strip()
    if r_str == "SENS": return "Sensitive"
    if "TB MR" in r_str or "TB R" in r_str: return "Resistant (Any)"
    if r_str == "AND": return "Not Evaluated"
    return "Not Evaluated"
itt_cohort["resistance_clean"] = itt_cohort["resistance"].apply(parse_res)

# Treatment duration for the index episode
itt_cohort["tx_days"] = (itt_cohort["end_date"] - itt_cohort["best_start"]).dt.days
def tx_grp(d):
    if d <= 60: return "< 2 months"
    if d <= 120: return "2 to <4 months"
    return "≥ 4 months"
itt_cohort["tx_month_grp"] = itt_cohort["tx_days"].apply(tx_grp)

# Save Final Cohort with well-organized time variables
keep_cols = ["sinan_clean", "age_tb", "age_group", "sex", "race_clean", "edu_clean", "hiv_aids", 
             "incarcerated", "homelessness", "dot_status", "alcohol", "drug_use", "mental_health", 
             "clinical_clean", "diagnosis_setting", "lab_confirmed_stat", "diabetes", "tobacco_use", 
             "hosp_admission", "other_immuno_condition", "bac1_clean", "sputum_culture_clean", "resistance_clean",
             "itt_group", "best_start", "end_date", "death_date", "tx_month_grp",
             "time_rn", "event_rn", "time_d", "event_d", "time_rn_tx", "time_d_tx"]

# re-organize for better readability (Time variables at the end)
itt_cohort[keep_cols].to_csv(OUT_CSV, index=False)
print(f"Saved optimized final ITT cohort to {OUT_CSV}")

# Build flowchart rows (step, n_remaining, n_excluded)
flow_rows = []
prev_n = None
for label, n in attrition:
    excluded = "" if prev_n is None else prev_n - n
    flow_rows.append({"step": label, "n_remaining": n, "n_excluded": excluded})
    prev_n = n

# Save exclusion flowchart as CSV (authoritative — regenerated on every run)
pd.DataFrame(flow_rows).to_csv(OUT_FLOWCHART, index=False)
print(f"Saved exclusion flowchart to {OUT_FLOWCHART}")

# Save Inclusion Doc
doc = Document()
doc.add_heading("Table 1/2 Cohort: Inclusion/Exclusion Flow (Mirrored Strict ITT)", 0)
doc.add_paragraph("This cohort matches the user's original eligibility criteria: First 'Novo' abandonment per individual.")
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Step"
hdr[1].text = "N Remaining"
hdr[2].text = "Excluded"
for row_data in flow_rows:
    row = table.add_row().cells
    row[0].text = row_data["step"]
    row[1].text = f"{row_data['n_remaining']:,}"
    row[2].text = "-" if row_data["n_excluded"] == "" else f"{row_data['n_excluded']:,}"
doc.save(OUT_DOC)
print(f"Saved inclusion/exclusion docx to {OUT_DOC}")
